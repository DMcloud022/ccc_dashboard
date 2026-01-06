"""
AI-Powered Action Plan Report Module for DICT Complaint Dashboard

This module generates strategic action plans based on complaint data analysis.
It is designed to work seamlessly with dashboard.py and receives pre-processed data.

DATA ALIGNMENT WITH DASHBOARD:
- Receives data already processed by dashboard.py's prepare_data() function
- Data has already been filtered to exclude Resolution = "FLS"
- Column mappings have already been applied
- Dates have been parsed and validated
- Service provider names have been normalized
- Invalid/empty rows have been removed

KEY FUNCTIONS:
- get_top_issues(): Identifies top 5 complaint issues from cleaned data
- generate_ai_action_plan(): Uses Gemini AI to create strategic action plans
- export_to_pdf/word(): Generates formatted reports for download
- render_weekly_report(): Main UI rendering function

USAGE:
This module is imported by dashboard.py and called in the AI Action Plan tab.
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import vertexai
from vertexai.generative_models import GenerativeModel
import google.auth
import json
from datetime import datetime
import os
from dotenv import load_dotenv
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

def clear_ai_report_state():
    """Clear the generated AI report state to force regeneration"""
    # Clear all report types with coverage periods
    report_types = ["Total_All_Complaints", "PEMEDES_Complaints_Only", "NTC_Complaints_Only"]
    coverage_periods = ["Monthly", "Quarterly", "Semi-Annual", "Annual"]
    
    for report_type in report_types:
        for period in coverage_periods:
            report_key = f"{report_type}_{period}"
            if f'report_generated_{report_key}' in st.session_state:
                st.session_state[f'report_generated_{report_key}'] = False
            if f'weekly_action_plan_{report_key}' in st.session_state:
                del st.session_state[f'weekly_action_plan_{report_key}']
            if f'executive_summary_{report_key}' in st.session_state:
                del st.session_state[f'executive_summary_{report_key}']
            if f'edited_action_plan_{report_key}' in st.session_state:
                del st.session_state[f'edited_action_plan_{report_key}']
            if f'sp_breakdowns_{report_key}' in st.session_state:
                del st.session_state[f'sp_breakdowns_{report_key}']
            if f'cached_pdf_bytes_{report_key}' in st.session_state:
                del st.session_state[f'cached_pdf_bytes_{report_key}']
            if f'cached_word_bytes_{report_key}' in st.session_state:
                del st.session_state[f'cached_word_bytes_{report_key}']
            if f'cached_csv_string_{report_key}' in st.session_state:
                del st.session_state[f'cached_csv_string_{report_key}']
    
    # Also clear legacy keys for backward compatibility
    if 'report_generated' in st.session_state:
        st.session_state.report_generated = False
    if 'weekly_action_plan' in st.session_state:
        del st.session_state.weekly_action_plan
    if 'executive_summary' in st.session_state:
        del st.session_state.executive_summary
    if 'edited_action_plan' in st.session_state:
        del st.session_state.edited_action_plan
    if 'sp_breakdowns' in st.session_state:
        del st.session_state.sp_breakdowns

# PEMEDES Service Providers List (Copied from dashboard.py for consistency)
PEMEDES_PROVIDERS = [
    "2GO Express, Inc.",
    "Sky Express",
    "3PL Service Provider Philippines, Inc.",
    "AAI Logistics Cargo Express d.b.a. Black Arrow",
    "A-Best Express, Inc.",
    "ACE-REM Messengerial and General Services, Inc.",
    "Airfreight 2100, Inc. d.b.a. Air21",
    "Airspeed International Corp.",
    "Arnold Mabanglo Express (AMEX)",
    "ASAP Courier and Logistics Phils., Inc.",
    "CACS Express and Allied Services, Inc.",
    "Cavatas-MSI Xpress Services Inc.",
    "CBL Freight Forwarder & Courier Express Int'l., Inc.",
    "Chargen Messengerial and General Services",
    "CJ Transnational Logistics Philippines, Inc.",
    "Cloud Panda PH, Inc. (d.b.a.) \"Tok Tok\"",
    "COMET Labor Service Cooperative",
    "DAG Xpress Courier, Inc.",
    "Del Asia Express Delivery & General Services, Inc.",
    "Diar's Assistance, Inc.",
    "Doorbell Technologies, Inc.",
    "El Grande Messengerial Services, Inc.",
    "Electrobill, Inc.",
    "Entrego Express Corporation",
    "EXMER, Inc.",
    "Fastcargo Logistics Corporation",
    "Fastrak Services, Inc.",
    "Fastrust Services, Inc.",
    "FES Business Solutions, Inc.",
    "Flash Express (PH) Co. Ltd., Inc.",
    "Flying High Energy Express Services Corporation",
    "GML Cargo Forwarder & Courier Express Int'l., Inc.",
    "GO21, Inc.",
    "GRABEXPRESS, Inc.",
    "Herald Express, Inc.",
    "Information Express Services, Inc. (INFORMEX)",
    "ICS / Intertraffic Transport Corporation",
    "Intervolve Express Services, Inc.",
    "Jay Messengerial and Manpower Services",
    "JG Manpo Janitorial & Messengerial Services Contractor",
    "Johnny Air Cargo and Delivery Services, Inc.",
    "JRMT Resources Corporation",
    "JRS Business Corporation",
    "LBC Express Corporation",
    "Libcap Super Express Corp.",
    "M.M. Bacarisas Courier Services",
    "Mail Expert Messengerial and General Services, Inc.",
    "Mailouwyz Courier",
    "Mailworld Express Service International, Inc.",
    "Mega Mail Express and General Services, Inc.",
    "MEJBAS Services, Inc.",
    "Mesco Express Service Corp.",
    "Metro Courier, Inc.",
    "Metro Prideco Services Corporation",
    "MMSC Services Corporation",
    "MR Messengerial & General Services",
    "MSPB Courier Services",
    "MTEL Trading & Manpower Services",
    "NGC Enterprises / Nathaniel G Cruz",
    "(Nathaniel G. Cruz Express Service)",
    "Ocean Coast Shipping Corp.",
    "Oceanwave Services, Inc.",
    "Pelican Express, Inc. to Cliqnship",
    "PH Gobal Jet Express, Inc. d.b.a. J&T Express",
    "PPB Messengerial Services, Inc.",
    "PRC Courier and Maintenance Services",
    "Priority Handling Logistics, Inc.",
    "PRO2000 Services, Inc.",
    "Promark Corporation",
    "Pronto Express Distribution, Inc.",
    "Quadx, Inc. d.b.a. GoGo Xpress",
    "Qualitrans Courier and Manpower Services, Inc.",
    "R&H Messengerial & General Services",
    "RAF International Forwarding Philippines, Inc.",
    "Republic Courier Service, Inc.",
    "RGServe Manpower Services",
    "RML Courier Services",
    "Safefreight Services, Inc.",
    "San Gabriel General Messengerial Services and Sales, Inc.",
    "Securetrac, Inc.",
    "Silver Royal General Services",
    "Snappmile Logistics, Inc.",
    "Speedels Services, Inc.",
    "Speedex Courier and Forwarder, Inc.",
    "Speedworks Courier Services Corporation",
    "Spex International Courier Services",
    "SPX Philippines, Inc.",
    "St. Joseph LFS Industrial Corp.",
    "Suremail Courier Services, Inc.",
    "Telexpress, Inc.",
    "TNT Express Deliveries (Phils.), Inc.",
    "Top Dynamics, Inc.",
    "Topserve Worldwide Express, Inc.",
    "Triload Express Systems",
    "UPS-Delbros Transport, Inc.",
    "Virgo Messengerial Services, Inc.",
    "Wall Street Courier Services, Inc. d.b.a. Ninja Van",
    "Wide Wide World Express Corporation",
    "Worklink Services, Inc.",
    "Ximex Delivery Express, Inc.",
    "Xytron International, Inc.",
    "ZIP Business Services, Inc.",
    "SF Express",
    "Mober PH",
    "Litexpress",
    "Lalamove",
    "YTO Express",
    "LEX PH",
    # Common variations and shortened names
    "2GO Express",
    "2GO",
    "J&T Express",
    "J&T",
    "LBC Express",
    "LBC",
    "Ninja Van",
    "Flash Express",
    "Grab Express",
    "GrabExpress",
    "Airspeed",
    "Air21",
    "Black Arrow",
    "GoGo Xpress",
    "SPX",
]

# NTC Service Providers List (Copied from dashboard.py for consistency)
NTC_PROVIDERS = [
    "PLDT",
    "Converge",
    "Globe",
    "Smart",
    "DITO",
    "Sky Cable",
    "Cignal",
    "Eastern",
    "DITO Telecommunity",
    "Globe Telecom",
    "Smart Communications",
    "PLDT Inc.",
    "Converge ICT Solutions",
    "Sky Fiber",
    "Royal Cable",
    "Radius Telecoms",
    "PT&T",
    "Now Telecom"
]

def is_ntc_complaint(agency):
    """Check if a complaint belongs to NTC (National Telecommunications Commission)"""
    if pd.isna(agency) or agency == '':
        return False

    # Handle non-string types
    if not isinstance(agency, str):
        agency = str(agency)

    agency_str = agency.strip()

    # Return False for empty strings after stripping
    if not agency_str:
        return False

    agency_lower = agency_str.lower()

    # Check for NTC (case-insensitive)
    # Match whole word to avoid false positives
    return 'ntc' in agency_lower.split() or agency_lower == 'ntc' or 'ntc' in agency_lower

def is_ntc_provider(service_provider):
    """Check if a service provider is an NTC provider"""
    if pd.isna(service_provider) or service_provider == '':
        return False

    # Handle non-string types
    if not isinstance(service_provider, str):
        service_provider = str(service_provider)

    service_provider_str = service_provider.strip()

    # Return False for empty strings after stripping
    if not service_provider_str:
        return False

    service_provider_lower = service_provider_str.lower()

    for ntc_sp in NTC_PROVIDERS:
        ntc_sp_lower = ntc_sp.lower()
        
        # Check for exact match or containment
        if ntc_sp_lower == service_provider_lower:
            return True
        if ntc_sp_lower in service_provider_lower:
            return True
        if service_provider_lower in ntc_sp_lower and len(service_provider_lower) > 3:
            return True
            
    return False

def is_pemedes_provider(service_provider):
    """Check if a service provider is a PEMEDES provider with robust matching"""
    if pd.isna(service_provider) or service_provider == '':
        return False

    # Handle non-string types
    if not isinstance(service_provider, str):
        service_provider = str(service_provider)

    service_provider_str = service_provider.strip()

    # Return False for empty strings after stripping
    if not service_provider_str:
        return False

    service_provider_lower = service_provider_str.lower()

    # Strategy 1: Exact match (case-insensitive)
    for pemedes_sp in PEMEDES_PROVIDERS:
        if service_provider_lower == pemedes_sp.lower():
            return True

    # Strategy 2 & 3: Containment checks (bidirectional)
    for pemedes_sp in PEMEDES_PROVIDERS:
        pemedes_sp_lower = pemedes_sp.lower()

        # Check if PEMEDES provider name is contained in the data
        if pemedes_sp_lower in service_provider_lower:
            # Additional safeguard: must match at least 5 characters or be a known short name
            if len(pemedes_sp_lower) >= 5 or pemedes_sp_lower in ['lbc', '2go', 'j&t', 'spx', 'air21']:
                return True

        # Check if data is contained in PEMEDES provider name (reverse check)
        if service_provider_lower in pemedes_sp_lower:
            # Additional safeguard: data must be at least 3 characters
            if len(service_provider_lower) >= 3:
                return True

    # Strategy 4: Multi-word name matching for complex cases
    for pemedes_sp in PEMEDES_PROVIDERS:
        pemedes_sp_lower = pemedes_sp.lower()

        # Only apply to multi-word names
        if ' ' in pemedes_sp_lower and ' ' in service_provider_lower:
            words_in_data = set(service_provider_lower.split())
            words_in_pemedes = set(pemedes_sp_lower.split())

            # Remove common words that don't help identify the provider
            common_words = {'inc', 'corp', 'corporation', 'express', 'services', 'service',
                          'courier', 'delivery', 'logistics', 'international', 'phils',
                          'philippines', 'ltd', 'co', 'and', 'the'}
            words_in_data = words_in_data - common_words
            words_in_pemedes = words_in_pemedes - common_words

            # If there's significant word overlap (at least 2 meaningful words)
            if len(words_in_pemedes) >= 2 and len(words_in_data.intersection(words_in_pemedes)) >= 2:
                return True

    return False

# DICT ORGANIZATIONAL STRUCTURE AND ISSUE MAPPING
# This mapping ensures accurate assignment of issues to the correct units and agencies

DICT_UNIT_MAPPING = {
    # 1. DELIVERY UNITS (Within DICT - Internal Services)
    "GDTB": {
        "name": "Government Digital Transformation Bureau",
        "keywords": ["egov", "pbh", "philippine business hub", "government services", "digital transformation"],
        "service_providers": ["GDTB", "eGOV", "NGP"]
    },
    "FPIAP": {
        "name": "Free Public Internet Access Program",
        "keywords": ["free wifi", "free wi-fi", "public internet", "community wifi", "pisonet"],
        "service_providers": ["FPIAP"]
    },
    "ILCDB": {
        "name": "ICT Literacy and Competency Development Bureau",
        "keywords": ["training", "upskilling", "certification", "certificate", "course", "learning", "workshop"],
        "service_providers": ["ILCDB"]
    },
    "AS": {
        "name": "Administrative Service",
        "keywords": ["hr concern", "human resource", "personnel", "employment", "recruitment"],
        "service_providers": ["AS", "HRMD", "HRDD", "GSD"]
    },
    "IMB": {
        "name": "Infrastructure Management Bureau",
        "keywords": ["cloud hosting", "web hosting", "government online", "gosd", "data center"],
        "service_providers": ["IMB", "GOSD", "GWHS", "NGDC"]
    },
    "CSB": {
        "name": "Cybersecurity Bureau",
        "keywords": ["digital certificate", "pki", "certificate authority", "encryption"],
        "service_providers": ["CSB", "PNPKI"]
    },
    "PRD": {
        "name": "Postal Regulation Division",
        "keywords": ["delivery concern", "courier", "logistics", "shipping", "parcel", "package"],
        "service_providers": ["LBC", "Ninja Van", "GO21", "GoGo Xpress", "Flash Express", "J&T", "J&T Express", "2GO", "SPX", "Lalamove"]
    },
    "ROCS": {
        "name": "Regional Operations and Coordination Service",
        "keywords": ["regional office", "regional concern", "regional service"],
        "service_providers": ["ROCS", "Regional Office"]
    },

    # 2. ATTACHED AGENCIES (Separate agencies under DICT supervision)
    "NTC": {
        "name": "National Telecommunications Commission",
        "keywords": ["internet", "telco", "disconnection", "slow connection", "network", "bandwidth", "fiber",
                    "broadband", "mobile data", "signal", "coverage", "technical service", "installation",
                    "unsolicited sms", "spam call", "telecom refund", "billing issue"],
        "service_providers": ["PLDT", "Converge", "Globe", "Smart", "DITO", "Sky Cable", "Cignal", "Eastern"]
    },
    "CICC": {
        "name": "Cybercrime Investigation and Coordinating Center",
        "keywords": ["cybercrime", "hacking", "phishing", "scam", "fraud", "identity theft", "online fraud",
                    "cyber attack", "data breach", "malware", "ransomware"],
        "service_providers": ["CICC"]
    },

    # 3. OTHER AGENCIES (External government bodies)
    "SEC": {
        "name": "Securities and Exchange Commission",
        "keywords": ["harassment", "online lending", "collection", "loan app", "lending app"],
        "service_providers": ["SEC"]
    },
    "DTI": {
        "name": "Department of Trade and Industry",
        "keywords": ["e-commerce", "online shopping", "consumer protection", "refund", "compensation",
                    "product quality", "false advertising", "lazada", "shopee"],
        "service_providers": ["DTI"]
    },
    "DOH": {
        "name": "Department of Health",
        "keywords": ["health", "vax", "vaccine", "vaccination", "health cert", "vax cert",
                    "immunization", "covid", "vaccination certificate"],
        "service_providers": ["DOH"]
    }
}

# Categorize units by organization type
DELIVERY_UNITS = ["GDTB", "FPIAP", "ILCDB", "AS", "IMB", "CSB", "PRD", "ROCS"]
ATTACHED_AGENCIES = ["NTC", "CICC"]
OTHER_AGENCIES = ["SEC", "DTI", "DOH"]

# Units that require service provider breakdown in reports
UNITS_REQUIRING_SP_BREAKDOWN = {
    "PRD": "Delivery Concerns",  # Show courier breakdown
    "NTC": "Telco/Internet Issues"  # Show ISP/telco breakdown
}

# Custom CSS for improved UI - Aligned with dashboard design
AI_REPORT_CSS = """
<style>
/* Import Google Fonts - Match Dashboard */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

/* AI Report Container */
.ai-report-container {
    background: white;
    border-radius: 8px;
    padding: 1.5rem;
    border: 1px solid #e5e7eb;
    margin-bottom: 0.75rem;
    font-family: 'Inter', sans-serif;
}

/* Report Header */
.report-header {
    text-align: center;
    padding-bottom: 1rem;
    border-bottom: 2px solid #3b82f6;
    margin-bottom: 1.5rem;
}

.report-title {
    font-size: 1.75rem;
    font-weight: 700;
    color: #1f2937;
    margin-bottom: 0.5rem;
    font-family: 'Inter', sans-serif;
}

.report-subtitle {
    font-size: 0.95rem;
    color: #6b7280;
    font-weight: 500;
}

/* Info Card */
.info-card {
    background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
    border-left: 4px solid #3b82f6;
    padding: 0.875rem 1rem;
    margin: 0.75rem 0;
    border-radius: 6px;
    font-size: 0.9rem;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
}

.info-card strong {
    color: #1f2937;
    font-weight: 600;
}

/* Responsive Table Container */
.responsive-table-container {
    width: 100%;
    overflow-x: auto;
    -webkit-overflow-scrolling: touch;
    border-radius: 8px;
    border: 1px solid #e5e7eb;
    margin: 1rem 0;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
}

/* Table Styling */
.action-plan-table {
    width: 100%;
    min-width: 800px;
    border-collapse: collapse;
    font-size: 0.9rem;
    font-family: 'Inter', sans-serif;
}

.action-plan-table th {
    background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
    color: white;
    padding: 0.875rem 1rem;
    text-align: left;
    font-weight: 600;
    font-size: 0.875rem;
    text-transform: uppercase;
    letter-spacing: 0.025em;
    position: sticky;
    top: 0;
    z-index: 10;
}

.action-plan-table td {
    padding: 0.875rem 1rem;
    border-bottom: 1px solid #e5e7eb;
    vertical-align: top;
    color: #374151;
    line-height: 1.5;
}

.action-plan-table tbody tr:hover {
    background-color: #f9fafb;
    transition: background-color 0.15s ease;
}

.action-plan-table tbody tr:last-child td {
    border-bottom: none;
}

/* Column Widths */
.col-issue {
    width: 25%;
    min-width: 180px;
    font-weight: 600;
    color: #1f2937;
}

.col-action {
    width: 35%;
    min-width: 260px;
}

.col-unit {
    width: 15%;
    min-width: 120px;
}

.col-remarks {
    width: 15%;
    min-width: 120px;
}

.col-resolution {
    width: 10%;
    min-width: 100px;
    text-align: center;
}

/* Mobile Responsiveness */
@media (max-width: 768px) {
    .ai-report-container {
        padding: 1rem;
    }

    .report-title {
        font-size: 1.5rem;
    }

    .report-subtitle {
        font-size: 0.85rem;
    }

    .action-plan-table {
        font-size: 0.85rem;
    }

    .action-plan-table th,
    .action-plan-table td {
        padding: 0.65rem 0.5rem;
    }

    .info-card {
        font-size: 0.85rem;
        padding: 0.75rem 0.875rem;
    }
}

/* Streamlit Button Overrides for AI Report */
div[data-testid="stButton"] > button {
    border-radius: 6px;
    font-weight: 600;
    font-size: 0.9rem;
    padding: 0.5rem 1.25rem;
    transition: all 0.2s ease;
    font-family: 'Inter', sans-serif;
}

div[data-testid="stButton"] > button[kind="primary"] {
    background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
    border: none;
    box-shadow: 0 2px 4px rgba(59, 130, 246, 0.2);
}

div[data-testid="stButton"] > button[kind="primary"]:hover {
    background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%);
    box-shadow: 0 4px 8px rgba(59, 130, 246, 0.3);
    transform: translateY(-1px);
}

div[data-testid="stDownloadButton"] > button {
    border-radius: 6px;
    font-weight: 600;
    font-size: 0.875rem;
    padding: 0.5rem 1rem;
    border: 1px solid #e5e7eb;
    background: white;
    color: #374151;
    transition: all 0.2s ease;
    font-family: 'Inter', sans-serif;
}

div[data-testid="stDownloadButton"] > button:hover {
    background: #f9fafb;
    border-color: #3b82f6;
    color: #3b82f6;
    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
    transform: translateY(-1px);
}

/* Data Editor Styling - Robust & Responsive */
div[data-testid="stDataFrame"] {
    width: 100%;
}

/* Ensure headers wrap properly */
div[data-testid="stDataFrame"] th {
    white-space: normal !important;
    vertical-align: top !important;
    padding: 8px !important;
    line-height: 1.4 !important;
}

/* Ensure cell content wraps and aligns properly */
div[data-testid="stDataFrame"] td {
    vertical-align: top !important;
}

/* Target the cell content div for wrapping */
div[data-testid="stDataFrame"] td div {
    white-space: pre-wrap !important;
    word-wrap: break-word !important;
    overflow-wrap: break-word !important;
    line-height: 1.5 !important;
    max-height: none !important;
}

/* Fix for Glide Data Grid (if used by Streamlit version) */
.glide-data-editor {
    font-family: 'Inter', sans-serif;
}
</style>
"""

def categorize_issue_to_unit(issue_name, issue_type="Category"):
    """
    Intelligently categorize an issue to the appropriate DICT unit or agency

    Args:
        issue_name: The name of the issue (from Complaint Category or Nature)
        issue_type: "Category" or "Nature"

    Returns:
        Tuple of (unit_code, unit_full_name, organization_type)
    """
    # Validate input
    if issue_name is None or issue_name == '':
        return "CICC", "Cybersecurity Investigation and Coordinating Center", "Attached Agency"

    issue_lower = str(issue_name).lower()

    # Score each unit based on keyword matches
    unit_scores = {}

    for unit_code, unit_info in DICT_UNIT_MAPPING.items():
        score = 0

        # Check keyword matches
        for keyword in unit_info["keywords"]:
            if keyword.lower() in issue_lower:
                score += 2  # Keyword match is worth 2 points

        # Check service provider matches (if available in the issue name)
        for provider in unit_info["service_providers"]:
            if provider.lower() in issue_lower:
                score += 3  # Service provider match is worth 3 points

        if score > 0:
            unit_scores[unit_code] = score

    # Return the unit with highest score
    if unit_scores:
        best_unit = max(unit_scores, key=unit_scores.get)

        # Determine organization type
        if best_unit in DELIVERY_UNITS:
            org_type = "Delivery Unit (DICT Internal)"
        elif best_unit in ATTACHED_AGENCIES:
            org_type = "Attached Agency"
        elif best_unit in OTHER_AGENCIES:
            org_type = "External Agency"
        else:
            org_type = "DICT"

        return best_unit, DICT_UNIT_MAPPING[best_unit]["name"], org_type

    # Fallback: categorize based on common patterns
    if any(word in issue_lower for word in ["internet", "connection", "network", "telco", "broadband"]):
        return "NTC", DICT_UNIT_MAPPING["NTC"]["name"], "Attached Agency"
    elif any(word in issue_lower for word in ["delivery", "courier", "shipping", "parcel"]):
        return "PRD", DICT_UNIT_MAPPING["PRD"]["name"], "Delivery Unit (DICT Internal)"
    elif any(word in issue_lower for word in ["cybercrime", "scam", "fraud", "hacking"]):
        return "CICC", DICT_UNIT_MAPPING["CICC"]["name"], "Attached Agency"
    elif any(word in issue_lower for word in ["ecommerce", "e-commerce", "shopping", "consumer"]):
        return "DTI", DICT_UNIT_MAPPING["DTI"]["name"], "External Agency"

    # Default fallback for unmatched issues
    return "CICC", "Cybersecurity Investigation and Coordinating Center", "Attached Agency"

@st.cache_resource(show_spinner=False)
def init_vertex_ai():
    """Initialize Vertex AI with error handling"""
    try:
        # Check if vertexai is properly imported
        if not hasattr(vertexai, 'init'):
            return False, "Vertex AI module not properly loaded. Please check your installation."

        credentials, project_id = google.auth.default()

        if not project_id:
            return False, "Project ID not found. Please set GOOGLE_CLOUD_PROJECT environment variable."

        vertexai.init(project=project_id, location="asia-southeast1")
        return True, None
    except ImportError as e:
        return False, f"Import error: {str(e)}. Please install required packages: pip install google-cloud-aiplatform"
    except Exception as e:
        return False, str(e)

def get_service_provider_breakdown(df, issue_name, issue_type):
    """
    Get service provider breakdown for a specific issue

    Args:
        df: The full complaint dataframe
        issue_name: The name of the issue (e.g., "Delivery Concerns (SP)" or normalized nature)
        issue_type: "Category" or "Nature"

    Returns:
        List of dicts with service provider counts and percentages
    """
    # Enhanced validation
    if df is None or df.empty:
        return []

    if 'Service Providers' not in df.columns:
        return []

    # Filter dataframe to only complaints matching this issue
    column_name = 'Complaint Category' if issue_type == "Category" else 'Complaint Nature'

    if column_name not in df.columns:
        return []

    try:
        if issue_type == "Category":
            # For categories, use exact match
            issue_complaints = df[df[column_name] == issue_name]
        else:
            # For nature, check both original and normalized matches
            # First try exact match
            issue_complaints = df[df[column_name] == issue_name]
            
            # If no exact match and this looks like a normalized name, 
            # find complaints that would normalize to this name
            if len(issue_complaints) == 0:
                mask = df[column_name].apply(lambda x: normalize_complaint_text(x) == issue_name)
                issue_complaints = df[mask]

        if len(issue_complaints) == 0:
            return []

        # Get service provider counts, but filter out inappropriate providers based on issue type
        sp_data = issue_complaints['Service Providers'].dropna()
        
        # For PEMEDES (Delivery Concerns), exclude NTC providers that might be miscategorized
        if issue_name == "Delivery Concerns (SP)" or "delivery" in issue_name.lower():
            sp_data = sp_data[sp_data.apply(lambda x: not is_ntc_provider(x))]
        
        # For NTC (Telco Issues), exclude PEMEDES providers that might be miscategorized
        elif issue_name == "Telco Internet Issues" or "telco" in issue_name.lower() or "internet" in issue_name.lower():
            sp_data = sp_data[sp_data.apply(lambda x: not is_pemedes_provider(x))]
        
        sp_counts = sp_data.value_counts()
    except Exception as e:
        # Silently handle any filtering errors
        return []

    if len(sp_counts) == 0:
        return []

    total_with_sp = sp_counts.sum()

    # Build breakdown list
    breakdown = []
    for provider, count in sp_counts.items():
        if provider and str(provider).strip():  # Skip empty values
            percentage = (count / total_with_sp * 100) if total_with_sp > 0 else 0
            breakdown.append({
                "provider": str(provider),
                "count": int(count),
                "percentage": round(percentage, 1)
            })

    # Sort by count descending and limit to top 5
    breakdown.sort(key=lambda x: x['count'], reverse=True)

    return breakdown[:5]  # Return only top 5 service providers

def normalize_complaint_text(text):
    """Normalize complaint text to handle similar descriptions"""
    if pd.isna(text) or text == '':
        return text
    
    text_str = str(text).strip().lower()
    
    # Common normalization patterns for delivery issues
    delivery_patterns = {
        'delayed/undelivered parcel': ['delayed parcel', 'undelivered parcel', 'delayed/ undelivered parcel', 
                                     'delayed /undelivered parcel', 'delayed undelivered parcel'],
        'mishandled parcel': ['mishandled parcel', 'damaged parcel', 'lost parcel'],
        'delivery concerns': ['delivery concern', 'delivery issue', 'delivery problem'],
        'billing issues': ['billing issue', 'billing problem', 'billing concern'],
        'internet disconnection': ['internet disconnection', 'internet disconnect', 'service disconnection'],
        'slow connection': ['slow internet', 'poor connection', 'slow connection'],
        'technical issues': ['technical issue', 'technical problem', 'technical concern']
    }
    
    # Check if text matches any pattern
    for normalized, variants in delivery_patterns.items():
        if any(variant in text_str for variant in variants):
            return normalized.title()  # Return normalized version with proper case
    
    # If no pattern matches, return original with proper case
    return str(text).strip().title()

@st.cache_data
def get_top_issues(df):
    """Extract top 5 issues based on Category and Nature with normalization

    This function aligns with dashboard.py's data structure and ensures
    we're analyzing the same cleaned data that's displayed in the dashboard.
    """
    if df is None or df.empty:
        return []

    issues = []

    try:
        # Analyze Categories (matching dashboard's approach)
        if 'Complaint Category' in df.columns:
            # Filter out NaN and empty values, matching dashboard behavior
            valid_categories = df['Complaint Category'].dropna()
            valid_categories = valid_categories[valid_categories != '']

            if len(valid_categories) > 0:
                top_cats = valid_categories.value_counts().head(5)
                for cat, count in top_cats.items():
                    issues.append({
                        "type": "Category",
                        "name": str(cat),  # Keep original category names
                        "count": int(count)
                    })

        # If we don't have enough categories, look at Nature with normalization
        if len(issues) < 5 and 'Complaint Nature' in df.columns:
            # Filter out NaN and empty values
            valid_nature = df['Complaint Nature'].dropna()
            valid_nature = valid_nature[valid_nature != '']

            if len(valid_nature) > 0:
                # Normalize nature descriptions to group similar ones
                normalized_nature = valid_nature.apply(normalize_complaint_text)
                
                # Count normalized values
                nature_counts = normalized_nature.value_counts()
                
                # Get top nature issues (excluding those already covered by categories)
                remaining_slots = 5 - len(issues)
                top_nature = nature_counts.head(remaining_slots)
                
                for nat, count in top_nature.items():
                    issues.append({
                        "type": "Nature",
                        "name": str(nat),  # Use normalized name
                        "count": int(count)
                    })
    except Exception as e:
        # Return empty list if any error occurs during analysis
        return []

    return issues[:5]

def generate_ai_action_plan(issues, df=None):
    """Generate action plan using Gemini

    This function creates strategic action plans based on the top complaint issues
    identified from the dashboard data (which has already been filtered and cleaned).
    Uses intelligent categorization to recommend the correct DICT unit or agency.

    Args:
        issues: List of top issues
        df: Optional dataframe for service provider analysis
    """
    if not issues or len(issues) == 0:
        return []

    # Validate that issues have required fields
    try:
        # First, categorize each issue to get recommended units and SP breakdown
        enriched_issues = []
        for issue in issues:
            if not isinstance(issue, dict) or 'name' not in issue or 'type' not in issue:
                continue  # Skip invalid issue entries

            unit_code, unit_name, org_type = categorize_issue_to_unit(issue['name'], issue['type'])

            # Get top service provider if applicable
            top_sp = None
            sp_count = 0
            sp_percentage = 0
            if df is not None and unit_code in UNITS_REQUIRING_SP_BREAKDOWN:
                sp_breakdown = get_service_provider_breakdown(df, issue['name'], issue['type'])
                if sp_breakdown and len(sp_breakdown) > 0:
                    top_sp = sp_breakdown[0]['provider']
                    sp_count = sp_breakdown[0]['count']
                    sp_percentage = sp_breakdown[0]['percentage']

            enriched_issues.append({
                **issue,
                "recommended_unit": unit_code,
                "recommended_unit_full": unit_name,
                "org_type": org_type,
                "top_service_provider": top_sp,
                "top_sp_count": sp_count,
                "top_sp_percentage": sp_percentage
            })

        # If no valid issues were enriched, return empty
        if len(enriched_issues) == 0:
            return []

    except Exception as e:
        # If enrichment fails, return empty list
        return []

    try:
        llm_model = os.getenv("LLM_MODEL", "gemini-1.5-flash-001")
        model = GenerativeModel(llm_model)

        system_prompt = os.getenv("SYSTEM_PROMPT", "You are a strategic analyst for the Department of Information and Communications Technology (DICT). Your role is to create actionable, specific, and measurable intervention plans to resolve citizen complaints.")

        # Build comprehensive unit guidelines with action plan templates
        unit_guidelines = """
DICT ORGANIZATIONAL STRUCTURE - UNIT ASSIGNMENT GUIDE WITH ACTION PLAN TEMPLATES:

1. DELIVERY UNITS (DICT Internal Services):
   - GDTB (Government Digital Transformation Bureau): eGov services, PBH, government digital transformation
     Template: "Conduct system audit of [specific service], implement fixes for [issue], and enhance user experience through [specific improvement]"

   - FPIAP (Free Public Internet Access Program): Free Wi-Fi, public internet access
     Template: "Deploy technical team to [location/issue area], restore/upgrade connectivity, and establish monitoring protocol"

   - ILCDB (ICT Literacy and Competency Development Bureau): Training, upskilling, certifications, courses
     Template: "Review [specific program], address [issue], streamline [process], and communicate timeline to affected participants"

   - AS (Administrative Service): HR concerns, personnel, recruitment
     Template: "Investigate [HR issue], implement corrective measures, and update policy/procedure to prevent recurrence"

   - IMB (Infrastructure Management Bureau): Cloud hosting, web hosting, government online services, data centers
     Template: "Conduct infrastructure assessment, resolve [technical issue], and implement redundancy/backup measures"

   - CSB (Cybersecurity Bureau): Digital certificates, PKI, encryption
     Template: "Fast-track certificate issuance/renewal process, resolve [specific issue], and establish expedited processing for backlog"

   - PRD (Postal Regulation Division): Delivery concerns, courier/logistics
     Template: "Escalate to [Top Service Provider] management, demand improved SLA compliance, establish penalty mechanism for delays, and explore alternative couriers"

   - ROCS (Regional Operations): Regional office concerns
     Template: "Coordinate with [specific region], deploy support team, resolve [issue], and strengthen regional coordination protocols"

2. ATTACHED AGENCIES (Under DICT supervision):
   - NTC (National Telecommunications Commission): Internet/telco issues, disconnections, slow connection, technical service
     Template: "Issue compliance directive to [Top Service Provider], mandate service restoration within [timeframe], impose penalties for SLA violations, and monitor resolution progress"

   - CICC (Cybersecurity Investigation and Coordinating Center): Cybercrime, hacking, phishing, scams, fraud
     Template: "Initiate investigation of [scam/fraud type], coordinate with law enforcement, issue public advisory, and pursue legal action against perpetrators"

3. OTHER AGENCIES (External partners):
   - SEC (Securities and Exchange Commission): Harassment, online lending, loan app collections
     Template: "Coordinate referral to SEC, provide complainant documentation support, and follow up on SEC enforcement action"

   - DTI (Department of Trade and Industry): E-commerce, consumer protection, retail refunds
     Template: "Refer to DTI Consumer Protection, facilitate mediation between parties, and support complaint resolution"
"""

        prompt = f"""
        {system_prompt}

        {unit_guidelines}

        Top Complaint Issues (pre-categorized with recommendations and service provider analysis):
        {json.dumps(enriched_issues, indent=2)}

        YOUR TASK: Create specific, actionable intervention plans for each issue.

        REQUIREMENTS FOR EACH ACTION PLAN:
        1. IDENTIFY ROOT CAUSE: What is the underlying problem causing this complaint?
        2. SPECIFIC ACTIONS: What concrete steps must be taken? (use the templates above as guides)
        3. TARGET SERVICE PROVIDER: If "top_service_provider" exists, name them specifically in the action plan
        4. MEASURABLE OUTCOME: What is the expected result?
        5. ACCOUNTABILITY: Who coordinates/implements? (use "recommended_unit" field)

        EXAMPLES OF GOOD ACTION PLANS:

        For NTC + Internet Issues + Top Provider "PLDT":
        "Issue compliance directive to PLDT requiring restoration of service within 48 hours for affected subscribers, impose administrative penalties for repeated SLA violations, establish weekly monitoring of complaint trends, and mandate quarterly service improvement reports."

        For PRD + Delivery Concerns + Top Provider "J&T Express":
        "Escalate to J&T Express regional management demanding immediate improvement in delivery timeframes, implement penalty mechanism for delays exceeding 5 days, require daily tracking updates for DICT shipments, and identify backup courier services to diversify risk."

        For NTC + Unsolicited SMS:
        "Direct all telecommunications providers to strengthen spam filtering mechanisms, issue show-cause orders to identified spam sources, coordinate with NBI Cybercrime Division for prosecution, and launch public awareness campaign on spam reporting procedures."

        For DTI + E-commerce refund issues:
        "Refer to DTI Consumer Protection Group, facilitate mediation between complainant and merchant, provide documentation support for filing formal complaints, and coordinate follow-up on resolution timeline."

        CRITICAL RULES:
        - DO NOT use generic language like "coordinate with" or "address concerns"
        - DO use specific verbs: "Issue directive", "Escalate to", "Implement", "Mandate", "Conduct audit", "Deploy team"
        - ALWAYS mention the top service provider by name if provided in the data
        - Include specific mechanisms (penalties, timelines, monitoring, reporting)
        - Use professional government language suitable for executive review
        - ALWAYS use the "recommended_unit" from the enriched_issues data
        - Keep action plans to 2-3 sentences maximum

        REMARKS FIELD INSTRUCTIONS:
        - Generate SPECIFIC, CONTEXTUAL remarks based on the actual data analysis for each issue
        - Include relevant metrics (complaint count, percentage, top provider if applicable)
        - Highlight priority level, urgency, or special considerations based on the actual issue
        - Examples of good remarks:
          * "Issue affects 234 users (45% of total complaints). Top provider J&T Express accounts for 67% of delivery issues."
          * "Critical priority - complaint volume increased 300% from previous period. Immediate intervention required."
          * "Recurring issue with PLDT services across multiple regions. Coordinate regulatory action with NTC."
          * "Limited to specific region. Monitor for pattern expansion before escalating intervention."
        - DO NOT use generic templates - base remarks on actual issue data provided above

        Return ONLY a valid JSON array with keys: "issue", "action_plan", "unit", "remarks".
        Do not include markdown formatting like ```json.
        """

        response = model.generate_content(prompt)

        # Clean response text to ensure valid JSON
        text = response.text.strip()
        if text.startswith("```json"):
            text = text[7:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()

        # Validate JSON before parsing
        if not text:
            raise ValueError("AI returned empty response")

        ai_plans = json.loads(text)

        # Validate that response is a list
        if not isinstance(ai_plans, list):
            raise ValueError("AI response is not a list of action plans")

        # Validate and correct unit assignments
        validated_plans = []
        for i, plan in enumerate(ai_plans):
            # Ensure plan is a dictionary
            if not isinstance(plan, dict):
                continue

            # Validate required fields exist
            if 'issue' not in plan or 'action_plan' not in plan or 'unit' not in plan:
                # Use enriched issue data if AI response is incomplete
                if i < len(enriched_issues):
                    plan = {
                        "issue": enriched_issues[i]["name"],
                        "action_plan": plan.get("action_plan", "Review and address complaints"),
                        "unit": enriched_issues[i]["recommended_unit"],
                        "remarks": ""
                    }

            # Use the pre-categorized unit if AI didn't assign correctly
            if plan.get("unit") not in DICT_UNIT_MAPPING:
                if i < len(enriched_issues):
                    plan["unit"] = enriched_issues[i]["recommended_unit"]

            # Ensure remarks field exists (AI should have generated this)
            if "remarks" not in plan or not plan["remarks"]:
                # Minimal fallback only if AI failed to generate remarks
                if i < len(enriched_issues):
                    issue_data = enriched_issues[i]
                    count = issue_data.get('count', 0)
                    top_sp = issue_data.get('top_service_provider', '')

                    # Create data-driven remark as minimal fallback
                    remark_parts = [f"Affects {count:,} complaints"]
                    if top_sp:
                        remark_parts.append(f"Top provider: {top_sp}")
                    plan["remarks"] = ". ".join(remark_parts) + ". Requires immediate attention."
                else:
                    plan["remarks"] = "Awaiting detailed analysis and implementation."

            validated_plans.append(plan)

        return validated_plans

    except Exception as e:
        st.error(f"AI Generation Error: {str(e)}")
        # Fallback: use pre-categorized units with specific action plans based on unit type and service provider
        fallback_plans = []
        for issue in enriched_issues:
            unit_code = issue['recommended_unit']
            unit_name = DICT_UNIT_MAPPING[unit_code]['name']
            top_sp = issue.get('top_service_provider')

            # Build specific action plans based on unit type
            if unit_code == "NTC":
                if top_sp:
                    action_plan = f"Issue compliance directive to {top_sp} requiring immediate service improvement, impose penalties for SLA violations, and establish monitoring mechanism for complaint resolution."
                else:
                    action_plan = f"Conduct investigation of telecommunications service quality issues, issue compliance directives to non-compliant providers, and enforce regulatory penalties where applicable."

            elif unit_code == "PRD":
                if top_sp:
                    action_plan = f"Escalate to {top_sp} management demanding improved delivery performance, implement penalty clauses for delays, and evaluate alternative courier services for future contracts."
                else:
                    action_plan = f"Review courier service provider contracts, enforce delivery SLA compliance, and establish performance monitoring system to prevent recurrence."

            elif unit_code == "CICC":
                action_plan = f"Initiate cybercrime investigation, coordinate with law enforcement agencies, issue public advisory on prevention measures, and pursue legal action against identified perpetrators."

            elif unit_code == "DTI":
                action_plan = f"Refer cases to DTI Consumer Protection Group, facilitate merchant-consumer mediation, provide complainants with documentation support, and coordinate follow-up on resolution timeline."

            elif unit_code == "SEC":
                action_plan = f"Coordinate referral to SEC Enforcement Department, assist complainants in filing formal complaints, and monitor SEC's regulatory action against violators."

            elif unit_code == "FPIAP":
                action_plan = f"Deploy technical team to assess connectivity issues, restore or upgrade affected infrastructure, and implement preventive monitoring system."

            elif unit_code == "GDTB":
                action_plan = f"Conduct comprehensive system audit, implement technical fixes for identified issues, and enhance user interface based on feedback analysis."

            elif unit_code == "ILCDB":
                action_plan = f"Review program implementation processes, address identified gaps in service delivery, streamline enrollment/certification procedures, and communicate updated timelines to participants."

            elif unit_code == "IMB":
                action_plan = f"Conduct infrastructure assessment, resolve technical service disruptions, implement system redundancy, and establish improved backup protocols."

            elif unit_code == "CSB":
                action_plan = f"Expedite digital certificate processing, address backlog in certificate issuance, and establish fast-track mechanism for urgent requests."

            else:
                # Generic fallback for any other unit
                action_plan = f"Conduct thorough investigation of complaints, implement corrective measures to address root causes, and establish monitoring system to prevent recurrence."

            # Generate minimal data-driven remarks for fallback (AI failure scenario)
            count = issue.get('count', 0)
            remark_parts = [f"Affects {count:,} complaints"]
            if top_sp:
                remark_parts.append(f"Top provider: {top_sp} ({sp_percentage:.1f}%)")
            remarks = ". ".join(remark_parts) + ". Requires prompt action."

            fallback_plans.append({
                "issue": issue['name'],
                "action_plan": action_plan,
                "unit": unit_code,
                "remarks": remarks
            })

        return fallback_plans

def generate_executive_summary(plans_data):
    """Generate an executive summary using AI based on the action plans"""
    try:
        llm_model = os.getenv("LLM_MODEL", "gemini-1.5-flash-001")
        model = GenerativeModel(llm_model)
        
        prompt = f"""
        You are a senior strategic analyst for the DICT. Based on the following Action Plans, generate a professional Executive Summary.
        
        Action Plans:
        {json.dumps(plans_data, indent=2)}
        
        Requirements:
        1. Write a main paragraph summarizing the overall situation (total complaints, top critical issues).
        2. Write separate short paragraphs for each Organization Type present in the data (e.g., "DICT Delivery Units", "Attached Agencies", "External Agencies").
        3. For each organization type, summarize the key issues and the planned interventions.
        4. Keep it concise, professional, and action-oriented.
        5. Do not use bullet points, use paragraphs.
        
        Output Format:
        Return a JSON object with the following keys:
        - "main_summary": The overall summary paragraph.
        - "org_summaries": A dictionary where keys are Organization Types (e.g. "Delivery Unit (DICT Internal)", "Attached Agency", "External Agency") and values are the summary paragraphs.
        """
        
        response = model.generate_content(prompt)
        text = response.text.strip()
        if text.startswith("```json"):
            text = text[7:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()
        
        return json.loads(text)
    except Exception as e:
        return {
            "main_summary": "Summary generation unavailable.",
            "org_summaries": {}
        }

def export_to_pdf(plans_df, top_issues, sp_breakdowns=None, dict_unit_counts=None, executive_summary=None, metrics=None, report_type="Total"):
    """Generate PDF report with service provider breakdowns

    Args:
        plans_df: DataFrame of action plans
        top_issues: List of top issues
        sp_breakdowns: Optional list of service provider breakdown data
        dict_unit_counts: Optional series of DICT unit counts
        executive_summary: Optional dictionary containing executive summary data
        metrics: Optional dictionary containing key metrics (Total, NTC, PEMEDES)
    """
    from reportlab.lib.pagesizes import landscape, A4

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), topMargin=0.5*inch, bottomMargin=0.5*inch,
                           leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )

    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#6b7280'),
        spaceAfter=20,
        alignment=TA_CENTER
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )

    # Title
    elements.append(Paragraph("DICT AI Action Plan Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", subtitle_style))
    elements.append(Spacer(1, 0.2*inch))

    # Metrics Section - Show only relevant metrics per report type
    if metrics:
        if "PEMEDES" in report_type:
            # PEMEDES Report - Only delivery metrics
            metrics_data = [[
                f"Delivery Complaints\n{metrics.get('total', 0):,}",
                f"Report Period\n{((metrics.get('end_date') - metrics.get('start_date')).days + 1) if metrics.get('end_date') and metrics.get('start_date') else 'N/A'} Days",
                f"Focus Area\nDelivery Services"
            ]]
            col_widths = [3.3*inch, 3.3*inch, 3.3*inch]
            backgrounds = [colors.HexColor('#dcfce7'), colors.HexColor('#f3f4f6'), colors.HexColor('#fef3c7')]
        elif "NTC" in report_type:
            # NTC Report - Only telecom metrics
            metrics_data = [[
                f"Telecom Complaints\n{metrics.get('total', 0):,}",
                f"Report Period\n{((metrics.get('end_date') - metrics.get('start_date')).days + 1) if metrics.get('end_date') and metrics.get('start_date') else 'N/A'} Days",
                f"Focus Area\nTelecommunications"
            ]]
            col_widths = [3.3*inch, 3.3*inch, 3.3*inch]
            backgrounds = [colors.HexColor('#e0f2fe'), colors.HexColor('#f3f4f6'), colors.HexColor('#fef3c7')]
        else:
            # Total Report - All metrics
            metrics_data = [[
                f"Total Complaints\n{metrics.get('total', 0):,}",
                f"Telecom Issues\n{metrics.get('ntc', 0):,}",
                f"Delivery Issues\n{metrics.get('pemedes', 0):,}"
            ]]
            col_widths = [3.3*inch, 3.3*inch, 3.3*inch]
            backgrounds = [colors.HexColor('#f3f4f6'), colors.HexColor('#e0f2fe'), colors.HexColor('#dcfce7')]
        
        metrics_table = Table(metrics_data, colWidths=col_widths)
        metrics_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1f2937')),
            ('BACKGROUND', (0, 0), (0, 0), backgrounds[0]),
            ('BACKGROUND', (1, 0), (1, 0), backgrounds[1]),
            ('BACKGROUND', (2, 0), (2, 0), backgrounds[2]),
            ('BOX', (0, 0), (-1, -1), 1, colors.white),
            ('TOPPADDING', (0, 0), (-1, -1), 15),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
        ]))
        elements.append(metrics_table)
        elements.append(Spacer(1, 0.3*inch))

    # Executive Summary
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#374151'),
        spaceAfter=6,
        leading=14
    )

    total_top5 = sum([issue['count'] for issue in top_issues])
    
    # Use dynamic summary if available
    if executive_summary and 'main_summary' in executive_summary:
        summary_text = f"<b>Executive Summary:</b><br/>{executive_summary['main_summary']}"
    else:
        summary_text = f"""
        <b>Executive Summary:</b><br/>
        This report identifies the top 5 priority complaint issues requiring immediate attention and provides strategic
        action plans for resolution. The analysis covers {total_top5:,} complaints from the top 5 issues, representing the most critical
        areas for intervention. Each issue has been assigned to the appropriate DICT unit or agency based on their
        mandate and expertise. Units are requested to review their assigned action plans, provide implementation remarks,
        and update resolution status as progress is made.
        """
    elements.append(Paragraph(summary_text, body_style))
    elements.append(Spacer(1, 0.2*inch))

    # Organization Summaries
    if executive_summary and 'org_summaries' in executive_summary and executive_summary['org_summaries']:
        elements.append(Paragraph("Summary by Organization Type", heading_style))
        
        # Map keys to display titles if needed, or just use keys
        categories = [
            ("DICT Delivery Units", ["Delivery Unit", "DICT Internal", "Delivery Unit (DICT Internal)"]),
            ("Attached Agencies", ["Attached Agency", "Attached Agencies"]),
            ("External Agencies", ["External Agency", "External Agencies"])
        ]
        
        # Display in specific order
        for title, keys in categories:
            summary_text = None
            for k, v in executive_summary['org_summaries'].items():
                if any(key_part.lower() in k.lower() for key_part in keys):
                    summary_text = v
                    break
            
            if summary_text:
                elements.append(Paragraph(f"<b>{title}</b>", body_style))
                elements.append(Paragraph(summary_text, body_style))
                elements.append(Spacer(1, 0.1*inch))
        
        # Catch any remaining summaries not in the categories
        for k, v in executive_summary['org_summaries'].items():
            matched = False
            for _, keys in categories:
                if any(key_part.lower() in k.lower() for key_part in keys):
                    matched = True
                    break
            if not matched:
                elements.append(Paragraph(f"<b>{k}</b>", body_style))
                elements.append(Paragraph(v, body_style))
                elements.append(Spacer(1, 0.1*inch))

        elements.append(Spacer(1, 0.1*inch))

    # Top Issues Summary
    elements.append(Paragraph("I. Top 5 Priority Issues", heading_style))

    issues_data = [['#', 'Issue', 'Source', 'Count']]
    for idx, issue in enumerate(top_issues, 1):
        issues_data.append([
            str(idx),
            Paragraph(issue['name'], styles['Normal']),
            issue['type'],  # Shows "Category" or "Nature"
            str(issue['count'])
        ])

    issues_table = Table(issues_data, colWidths=[0.5*inch, 3.5*inch, 1*inch, 0.8*inch])
    issues_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
    ]))

    elements.append(issues_table)

    # Add footnote for Source column
    footnote_style = ParagraphStyle(
        'Footnote',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#6b7280'),
        spaceAfter=6,
        leftIndent=0.5*inch
    )
    elements.append(Paragraph("Note: 'Source' indicates whether the issue is from 'Category' or 'Nature' field in complaint data.", footnote_style))
    elements.append(Spacer(1, 0.3*inch))

    # Action Plan Table
    elements.append(Paragraph("II. Strategic Action Plan Details", heading_style))

    # Add narrative explanation
    action_plan_narrative = f"""
    The following action plans have been developed for each priority issue. Each plan outlines specific interventions
    required and identifies the responsible DICT unit or agency. The Remarks and Resolution columns are provided for
    units to document their implementation progress, challenges encountered, and final resolution status.
    """
    elements.append(Paragraph(action_plan_narrative, body_style))
    elements.append(Spacer(1, 0.15*inch))

    # Define header style for table
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        textColor=colors.whitesmoke,
        alignment=TA_LEFT
    )

    # Use Paragraph for headers to ensure wrapping and prevent overlap
    plan_data = [[
        Paragraph('Issue', header_style),
        Paragraph('Action Plan', header_style),
        Paragraph('Assigned Unit', header_style),
        Paragraph('Remarks', header_style),
        Paragraph('Action Taken', header_style)
    ]]
    
    for _, row in plans_df.iterrows():
        # Use actual values from edited data (remarks and resolution may be edited)
        remarks_text = str(row.get('remarks', '')) if pd.notna(row.get('remarks', '')) else ''
        resolution_text = str(row.get('resolution', '')) if pd.notna(row.get('resolution', '')) else ''

        plan_data.append([
            Paragraph(str(row['issue']), styles['Normal']),
            Paragraph(str(row['action_plan']), styles['Normal']),
            Paragraph(str(row['unit']), styles['Normal']),
            Paragraph(remarks_text, styles['Normal']),
            Paragraph(resolution_text, styles['Normal'])
        ])

    # Landscape A4 width is about 11 inches, minus margins = ~10 inches available
    # Adjusted column widths to prevent overlap and improve readability
    # Issue: 1.5, Action Plan: 3.8, Unit: 1.0, Remarks: 2.5, Action Taken: 1.5 = 10.3 inches
    col_widths = [1.5*inch, 3.8*inch, 1.0*inch, 2.5*inch, 1.5*inch]
    
    plan_table = Table(plan_data, colWidths=col_widths)
    plan_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        # Remove direct font settings for header row as we use Paragraph now
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
    ]))

    elements.append(plan_table)

    # Add note about editable fields
    note_style = ParagraphStyle(
        'Note',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#6b7280'),
        spaceAfter=6,
        leftIndent=0.5*inch,
        italic=True
    )
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph("Note: This report includes any edits made to the Action Plan, Remarks, and Action Taken by the Unit fields before download.", note_style))

    # Service Provider Breakdown Section
    if sp_breakdowns and len(sp_breakdowns) > 0:
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph("III. Service Provider Analysis", heading_style))

        # Add narrative
        sp_narrative = """
        For issues related to delivery concerns and telecommunications, the following breakdown identifies specific
        service providers responsible for the majority of complaints. This granular analysis enables targeted
        interventions with individual providers to address service quality issues effectively. Each breakdown shows
        the top 5 service providers contributing to the issue, allowing focused engagement with the most problematic providers.
        """
        elements.append(Paragraph(sp_narrative, body_style))
        elements.append(Spacer(1, 0.15*inch))

        body_style_sp = ParagraphStyle(
            'BodySP',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#374151')
        )

        for sp_item in sp_breakdowns:
            # Issue header
            issue_header = f"{sp_item['issue']} ({sp_item['unit']}) - {sp_item['total_count']} total complaints"
            elements.append(Paragraph(issue_header, styles['Heading3']))
            elements.append(Spacer(1, 0.1*inch))

            # Add table-specific narrative
            num_providers = len(sp_item['breakdown'])
            top_provider_pct = sp_item['breakdown'][0]['percentage'] if sp_item['breakdown'] else 0

            table_narrative = f"""
            The table below presents the top {num_providers} service providers for this issue category.
            The leading provider accounts for {top_provider_pct:.1f}% of complaints in this category,
            indicating {"a concentrated issue requiring focused intervention" if top_provider_pct > 40 else "a distributed problem across multiple providers"}.
            Coordinating with these providers can directly address {sum([sp['percentage'] for sp in sp_item['breakdown']]):.1f}% of complaints in this category.
            """
            elements.append(Paragraph(table_narrative, body_style_sp))
            elements.append(Spacer(1, 0.1*inch))

            # SP breakdown table
            sp_data = [['Service Provider', 'Complaints', 'Percentage']]
            for sp in sp_item['breakdown']:
                sp_data.append([
                    sp['provider'],
                    str(sp['count']),
                    f"{sp['percentage']}%"
                ])

            sp_table = Table(sp_data, colWidths=[3*inch, 1.2*inch, 1.2*inch])
            sp_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1f2937')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))

            elements.append(sp_table)
            elements.append(Spacer(1, 0.15*inch))

            # Top provider analysis and recommendation
            if sp_item['breakdown']:
                top_sp = sp_item['breakdown'][0]

                # Generate actionable recommendation based on data
                if top_sp['percentage'] > 50:
                    recommendation = f"Immediate escalation to {top_sp['provider']} management is recommended as they represent the majority of issues."
                elif top_sp['percentage'] > 30:
                    recommendation = f"Priority engagement with {top_sp['provider']} while monitoring other providers is advised."
                else:
                    recommendation = f"A multi-provider approach is recommended given the distributed nature of complaints."

                analysis_text = f"""
                <b>Key Finding:</b> {top_sp['provider']} leads with {top_sp['count']} complaints ({top_sp['percentage']}%).
                <b>Recommended Action:</b> {recommendation}
                """
                elements.append(Paragraph(analysis_text, body_style_sp))
                elements.append(Spacer(1, 0.2*inch))

    # DICT Unit Breakdown
    if dict_unit_counts is not None and len(dict_unit_counts) > 0:
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph("IV. Complaints by DICT Unit", heading_style))
        
        # Add description
        top_unit = dict_unit_counts.index[0]
        top_count = dict_unit_counts.iloc[0]
        total_unit_complaints = dict_unit_counts.sum()
        desc_text = f"The table below details the distribution of complaints across different DICT units. {top_unit} received the highest volume with {top_count} complaints, accounting for {(top_count/total_unit_complaints*100):.1f}% of the top unit-attributed complaints."
        elements.append(Paragraph(desc_text, body_style))
        elements.append(Spacer(1, 0.15*inch))
        
        unit_data = [['DICT Unit', 'Count']]
        for unit, count in dict_unit_counts.items():
            unit_data.append([str(unit), str(count)])
            
        unit_table = Table(unit_data, colWidths=[4*inch, 1.5*inch])
        unit_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
        ]))
        elements.append(unit_table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

def export_to_word(plans_df, top_issues, sp_breakdowns=None, dict_unit_counts=None, executive_summary=None, metrics=None, report_type="Total"):
    """Generate Word document report with service provider breakdowns

    Args:
        plans_df: DataFrame of action plans
        top_issues: List of top issues
        sp_breakdowns: Optional list of service provider breakdown data
        dict_unit_counts: Optional series of DICT unit counts
        executive_summary: Optional dictionary containing executive summary data
        metrics: Optional dictionary containing key metrics (Total, NTC, PEMEDES)
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_cell_background(cell, fill_color):
        """Set cell background color"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), fill_color)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except:
            pass  # Silently fail if shading doesn't work

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('DICT AI Action Plan Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(31, 41, 55)

    # Subtitle
    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    # Metrics Section - Show only relevant metrics per report type
    if metrics:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = True
        
        row = table.rows[0]
        
        if "PEMEDES" in report_type:
            # PEMEDES Report - Only delivery metrics
            cell0 = row.cells[0]
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run(f"Delivery Complaints\n{metrics.get('total', 0):,}")
            run0.font.bold = True
            set_cell_background(cell0, "DCFCE7")
            
            cell1 = row.cells[1]
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            period_days = ((metrics.get('end_date') - metrics.get('start_date')).days + 1) if metrics.get('end_date') and metrics.get('start_date') else 'N/A'
            run1 = p1.add_run(f"Report Period\n{period_days} Days")
            run1.font.bold = True
            set_cell_background(cell1, "F3F4F6")
            
            cell2 = row.cells[2]
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(f"Focus Area\nDelivery Services")
            run2.font.bold = True
            set_cell_background(cell2, "FEF3C7")
            
        elif "NTC" in report_type:
            # NTC Report - Only telecom metrics
            cell0 = row.cells[0]
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run(f"Telecom Complaints\n{metrics.get('total', 0):,}")
            run0.font.bold = True
            set_cell_background(cell0, "E0F2FE")
            
            cell1 = row.cells[1]
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            period_days = ((metrics.get('end_date') - metrics.get('start_date')).days + 1) if metrics.get('end_date') and metrics.get('start_date') else 'N/A'
            run1 = p1.add_run(f"Report Period\n{period_days} Days")
            run1.font.bold = True
            set_cell_background(cell1, "F3F4F6")
            
            cell2 = row.cells[2]
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(f"Focus Area\nTelecommunications")
            run2.font.bold = True
            set_cell_background(cell2, "FEF3C7")
            
        else:
            # Total Report - All metrics
            cell0 = row.cells[0]
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run(f"Total Complaints\n{metrics.get('total', 0):,}")
            run0.font.bold = True
            set_cell_background(cell0, "F3F4F6")
            
            cell1 = row.cells[1]
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p1.add_run(f"Telecom Issues\n{metrics.get('ntc', 0):,}")
            run1.font.bold = True
            set_cell_background(cell1, "E0F2FE")
            
            cell2 = row.cells[2]
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(f"Delivery Issues\n{metrics.get('pemedes', 0):,}")
            run2.font.bold = True
            set_cell_background(cell2, "DCFCE7")
        
        doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    if executive_summary and 'main_summary' in executive_summary:
        doc.add_paragraph(executive_summary['main_summary'])
    else:
        doc.add_paragraph("This report identifies the top 5 priority complaint issues requiring immediate attention and provides strategic action plans for resolution.")
    
    doc.add_paragraph()

    # Organization Summaries
    if executive_summary and 'org_summaries' in executive_summary and executive_summary['org_summaries']:
        doc.add_heading('Summary by Organization Type', 2)
        
        categories = [
            ("DICT Delivery Units", ["Delivery Unit", "DICT Internal", "Delivery Unit (DICT Internal)"]),
            ("Attached Agencies", ["Attached Agency", "Attached Agencies"]),
            ("External Agencies", ["External Agency", "External Agencies"])
        ]
        
        for title, keys in categories:
            summary_text = None
            for k, v in executive_summary['org_summaries'].items():
                if any(key_part.lower() in k.lower() for key_part in keys):
                    summary_text = v
                    break
            
            if summary_text:
                p = doc.add_paragraph()
                runner = p.add_run(title)
                runner.bold = True
                doc.add_paragraph(summary_text)
                doc.add_paragraph()
        
        # Catch remaining
        for k, v in executive_summary['org_summaries'].items():
            matched = False
            for _, keys in categories:
                if any(key_part.lower() in k.lower() for key_part in keys):
                    matched = True
                    break
            if not matched:
                p = doc.add_paragraph()
                runner = p.add_run(k)
                runner.bold = True
                doc.add_paragraph(v)
                doc.add_paragraph()

    # Top Issues Section
    doc.add_heading('Top 5 Priority Issues', 1)

    issues_table = doc.add_table(rows=1, cols=4)
    issues_table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = issues_table.rows[0].cells
    header_cells[0].text = '#'
    header_cells[1].text = 'Issue'
    header_cells[2].text = 'Source'
    header_cells[3].text = 'Count'

    # Format header
    for cell in header_cells:
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '3B82F6')

    # Data rows
    for idx, issue in enumerate(top_issues, 1):
        row_cells = issues_table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = issue['name']
        row_cells[2].text = issue['type']  # Shows "Category" or "Nature"
        row_cells[3].text = str(issue['count'])

    # Add footnote for Source column
    note = doc.add_paragraph()
    note_run = note.add_run("Note: 'Source' indicates whether the issue is from 'Category' or 'Nature' field in complaint data.")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(107, 114, 128)
    note_run.italic = True

    doc.add_paragraph()

    # Action Plan Section
    doc.add_heading('Action Plan Details', 1)

    plan_table = doc.add_table(rows=1, cols=5)
    plan_table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = plan_table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Action Plan'
    header_cells[2].text = 'Assigned Unit'
    header_cells[3].text = 'Remarks'
    header_cells[4].text = 'Action Taken by the Unit'

    # Format header
    for cell in header_cells:
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '3B82F6')

    # Data rows
    for _, row in plans_df.iterrows():
        # Use actual values from edited data (remarks and resolution may be edited)
        remarks_text = str(row.get('remarks', '')) if pd.notna(row.get('remarks', '')) else ''
        resolution_text = str(row.get('resolution', '')) if pd.notna(row.get('resolution', '')) else ''

        row_cells = plan_table.add_row().cells
        row_cells[0].text = str(row['issue'])
        row_cells[1].text = str(row['action_plan'])
        row_cells[2].text = str(row['unit'])
        row_cells[3].text = remarks_text
        row_cells[4].text = resolution_text
        # Center align the resolution
        if row_cells[4].paragraphs:
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add note about editable fields
    note = doc.add_paragraph()
    note_run = note.add_run("Note: This report includes any edits made to the Action Plan, Remarks, and Action Taken by the Unit fields before download.")
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(107, 114, 128)
    note_run.italic = True

    # Service Provider Breakdown Section
    if sp_breakdowns and len(sp_breakdowns) > 0:
        doc.add_paragraph()
        doc.add_heading('Service Provider Analysis', 1)

        # Add section narrative
        section_intro = doc.add_paragraph()
        intro_text = (
            "For issues related to delivery concerns and telecommunications, the following breakdown identifies "
            "specific service providers responsible for the majority of complaints. This granular analysis enables "
            "targeted interventions with individual providers to address service quality issues effectively. "
            "Each breakdown shows the top 5 service providers contributing to the issue, allowing focused "
            "engagement with the most problematic providers."
        )
        intro_run = section_intro.add_run(intro_text)
        intro_run.font.size = Pt(10)
        intro_run.font.color.rgb = RGBColor(55, 65, 81)
        doc.add_paragraph()

        for sp_item in sp_breakdowns:
            # Issue subheading
            issue_heading = doc.add_heading(f"{sp_item['issue']} ({sp_item['unit']}) - {sp_item['total_count']} complaints", 2)

            # Add table-specific narrative
            num_providers = len(sp_item['breakdown'])
            top_provider_pct = sp_item['breakdown'][0]['percentage'] if sp_item['breakdown'] else 0

            table_narrative_para = doc.add_paragraph()
            table_narrative_text = (
                f"The table below presents the top {num_providers} service providers for this issue category. "
                f"The leading provider accounts for {top_provider_pct:.1f}% of complaints in this category, "
                f"indicating {'a concentrated issue requiring focused intervention' if top_provider_pct > 40 else 'a distributed problem across multiple providers'}. "
                f"Coordinating with these providers can directly address {sum([sp['percentage'] for sp in sp_item['breakdown']]):.1f}% of complaints in this category."
            )
            narrative_run = table_narrative_para.add_run(table_narrative_text)
            narrative_run.font.size = Pt(10)
            narrative_run.font.color.rgb = RGBColor(55, 65, 81)
            doc.add_paragraph()

            # SP breakdown table
            sp_table = doc.add_table(rows=1, cols=3)
            sp_table.style = 'Light Grid Accent 1'

            # Header
            sp_header_cells = sp_table.rows[0].cells
            sp_header_cells[0].text = 'Service Provider'
            sp_header_cells[1].text = 'Complaints'
            sp_header_cells[2].text = 'Percentage'

            for cell in sp_header_cells:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, 'F3F4F6')

            # Data rows
            for sp in sp_item['breakdown']:
                sp_row_cells = sp_table.add_row().cells
                sp_row_cells[0].text = sp['provider']
                sp_row_cells[1].text = str(sp['count'])
                sp_row_cells[2].text = f"{sp['percentage']}%"

            # Top provider analysis and recommendation
            if sp_item['breakdown']:
                top_sp = sp_item['breakdown'][0]

                # Generate actionable recommendation based on data
                if top_sp['percentage'] > 50:
                    recommendation = f"Immediate escalation to {top_sp['provider']} management is recommended as they represent the majority of issues."
                elif top_sp['percentage'] > 30:
                    recommendation = f"Priority engagement with {top_sp['provider']} while monitoring other providers is advised."
                else:
                    recommendation = f"A multi-provider approach is recommended given the distributed nature of complaints."

                # Key Finding
                finding_para = doc.add_paragraph()
                finding_label = finding_para.add_run("Key Finding: ")
                finding_label.font.bold = True
                finding_label.font.size = Pt(9)
                finding_label.font.color.rgb = RGBColor(31, 41, 55)

                finding_text = finding_para.add_run(f"{top_sp['provider']} leads with {top_sp['count']} complaints ({top_sp['percentage']}%).")
                finding_text.font.size = Pt(9)
                finding_text.font.color.rgb = RGBColor(55, 65, 81)

                # Recommended Action
                action_para = doc.add_paragraph()
                action_label = action_para.add_run("Recommended Action: ")
                action_label.font.bold = True
                action_label.font.size = Pt(9)
                action_label.font.color.rgb = RGBColor(31, 41, 55)

                action_text = action_para.add_run(recommendation)
                action_text.font.size = Pt(9)
                action_text.font.color.rgb = RGBColor(55, 65, 81)

            doc.add_paragraph()

    # DICT Unit Breakdown
    if dict_unit_counts is not None and len(dict_unit_counts) > 0:
        doc.add_paragraph()
        doc.add_heading('Complaints by DICT Unit', 1)
        
        # Add description
        top_unit = dict_unit_counts.index[0]
        top_count = dict_unit_counts.iloc[0]
        total_unit_complaints = dict_unit_counts.sum()
        desc_text = f"The table below details the distribution of complaints across different DICT units. {top_unit} received the highest volume with {top_count} complaints, accounting for {(top_count/total_unit_complaints*100):.1f}% of the top unit-attributed complaints."
        doc.add_paragraph(desc_text)
        doc.add_paragraph()
        
        unit_table = doc.add_table(rows=1, cols=2)
        unit_table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = unit_table.rows[0].cells
        header_cells[0].text = 'DICT Unit'
        header_cells[1].text = 'Count'
        
        for cell in header_cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, '3B82F6')
            
        # Data
        for unit, count in dict_unit_counts.items():
            row_cells = unit_table.add_row().cells
            row_cells[0].text = str(unit)
            row_cells[1].text = str(count)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_weekly_report(df, filter_year=None, filter_month=None):
    """Render the Weekly Report / Action Plan section with improved UI
    
    Args:
        df: DataFrame with complaint data
        filter_year: Selected year from dashboard filter (None means use all data)
        filter_month: Selected month from dashboard filter (0 or None means all months in year)
    """

    # Apply custom CSS
    st.markdown(AI_REPORT_CSS, unsafe_allow_html=True)

    # Header
    st.markdown("""
    <div class="report-header">
        <div class="report-title">DICT AI-Powered Action Plan Reports</div>
        <div class="report-subtitle">Strategic Analysis and Recommendations for Priority Complaint Issues</div>
    </div>
    """, unsafe_allow_html=True)

    # Check if dashboard date filter is active
    # Filter is active if: year is set AND year is not "All Years" 
    dashboard_filter_active = (filter_year is not None and 
                              filter_year != "All Years" and 
                              str(filter_year).strip() != "")
    
    # Report Type and Coverage Selection
    if dashboard_filter_active:
        # When dashboard filter is active, only show report type
        report_type = st.selectbox(
            "Report Type:",
            ["Total (All Complaints)", "PEMEDES Complaints Only", "NTC Complaints Only"],
            help="Select the specific report type to generate individual action plans for different stakeholders"
        )
        # Disable coverage period when dashboard filter is active
        coverage_period = None
        
        # Show clear message about what data is being used
        if filter_month and filter_month != 0:
            month_name = datetime(2020, filter_month, 1).strftime('%B')
            st.info(f"📅 **Report uses Dashboard filter:** {month_name} {filter_year} (Report Coverage disabled)")
        else:
            st.info(f"📅 **Report uses Dashboard filter:** Year {filter_year} (Report Coverage disabled)")
    else:
        # When no dashboard filter, show both report type and coverage period
        col_type, col_period = st.columns([2, 1])
        
        with col_type:
            report_type = st.selectbox(
                "Report Type:",
                ["Total (All Complaints)", "PEMEDES Complaints Only", "NTC Complaints Only"],
                help="Select the specific report type to generate individual action plans for different stakeholders"
            )
        
        with col_period:
            coverage_period = st.selectbox(
                "Report Coverage:",
                ["Monthly", "Quarterly", "Semi-Annual", "Annual"],
                index=0,  # Monthly is default
                help="Select the time period coverage for the report analysis"
            )
    
    st.markdown("---")

    if df is None or df.empty:
        st.info("No data available to generate report. Please load complaint data from the Dashboard tab first.")
        st.markdown("""
        **Instructions:**
        1. Navigate to the Dashboard tab
        2. Load complaint data using the sidebar options
        3. Return to this tab to generate the AI-powered action plan
        """)
        return

    # Apply date filter from dashboard if provided
    if dashboard_filter_active and 'Date Received' in df.columns:
        if filter_month and filter_month != 0:
            # Filter by both year and month
            df = df[(df['Date Received'].dt.year == filter_year) & 
                   (df['Date Received'].dt.month == filter_month)].copy()
        else:
            # Filter by year only
            df = df[df['Date Received'].dt.year == filter_year].copy()
        
        if df.empty:
            st.warning(f"⚠️ No data found for the selected period in the Dashboard.")
            return
    elif (filter_year is not None and 
          str(filter_year) == "All Years" and 
          filter_month and 
          filter_month != 0 and 
          'Date Received' in df.columns):
        # Filter by month across all years
        df = df[df['Date Received'].dt.month == filter_month].copy()
        
        if df.empty:
            st.warning(f"⚠️ No data found for the selected month across all years.")
            return

    # IMPORTANT: The df passed here is already prepared by dashboard.py's prepare_data() function
    # which handles:
    # - Column mapping and validation
    # - Date parsing
    # - Service provider normalization
    # - FLS resolution filtering (Resolution != "FLS")
    # - Invalid date removal
    # Therefore, we work with the data as-is without re-processing

    # Filter data based on report type
    if report_type == "PEMEDES Complaints Only":
        if 'Complaint Category' in df.columns:
            # Primary filter: Delivery Concerns category
            pemedes_mask = (df['Complaint Category'].astype(str).str.strip().str.upper() == "DELIVERY CONCERNS (SP)")
            df_filtered_category = df[pemedes_mask].copy()
            
            # Additional filter: Exclude NTC service providers that might be miscategorized
            if 'Service Providers' in df_filtered_category.columns:
                # Filter out any NTC providers (telecom companies) from PEMEDES report
                ntc_provider_mask = df_filtered_category['Service Providers'].apply(
                    lambda x: not is_ntc_provider(x) if pd.notna(x) else True
                )
                df_base = df_filtered_category[ntc_provider_mask].copy()
            else:
                df_base = df_filtered_category.copy()
            
            report_title_suffix = " - PEMEDES Delivery Concerns"
        else:
            st.error("Cannot filter PEMEDES complaints: 'Complaint Category' column not found")
            return
    elif report_type == "NTC Complaints Only":
        if 'Complaint Category' in df.columns:
            # Primary filter: Telco Internet Issues category
            ntc_mask = (df['Complaint Category'].astype(str).str.strip().str.upper() == "TELCO INTERNET ISSUES")
            df_filtered_category = df[ntc_mask].copy()
            
            # Additional filter: Exclude PEMEDES providers that might be miscategorized
            if 'Service Providers' in df_filtered_category.columns:
                # Filter out any PEMEDES providers (courier/delivery companies) from NTC report
                pemedes_provider_mask = df_filtered_category['Service Providers'].apply(
                    lambda x: not is_pemedes_provider(x) if pd.notna(x) else True
                )
                df_base = df_filtered_category[pemedes_provider_mask].copy()
            else:
                df_base = df_filtered_category.copy()
            
            report_title_suffix = " - NTC Telecommunications (Telco Internet Issues)"
        else:
            st.error("Cannot filter NTC complaints: 'Complaint Category' column not found")
            return
    else:  # Total (All Complaints)
        df_base = df.copy()
        report_title_suffix = " - All Complaints"

    # Dynamic Date Filtering based on Coverage Period (only if dashboard filter is NOT active)
    metrics = {}
    if 'Date Received' in df_base.columns:
        valid_dates = df_base['Date Received'].dropna()
        if len(valid_dates) > 0:
            max_date_avail = valid_dates.max()
            
            # Only apply coverage period filtering if dashboard filter is not active
            if coverage_period and not dashboard_filter_active:
                # Calculate date range based on coverage period
                if coverage_period == "Monthly":
                    start_date = max_date_avail - pd.DateOffset(months=1)
                elif coverage_period == "Quarterly":
                    start_date = max_date_avail - pd.DateOffset(months=3)
                elif coverage_period == "Semi-Annual":
                    start_date = max_date_avail - pd.DateOffset(months=6)
                elif coverage_period == "Annual":
                    start_date = max_date_avail - pd.DateOffset(months=12)
                else:
                    # Default to monthly if invalid selection
                    start_date = max_date_avail - pd.DateOffset(months=1)
                
                end_date = max_date_avail
                
                # Filter data
                mask = (df_base['Date Received'] >= start_date) & (df_base['Date Received'] <= end_date)
            else:
                # Dashboard filter is active - use all data from df_base (already filtered)
                # Set date range to the actual min/max of the filtered data
                start_date = valid_dates.min()
                end_date = max_date_avail
                mask = pd.Series([True] * len(df_base), index=df_base.index)
            df_filtered = df_base[mask].copy()
            
            # Calculate ALL-TIME totals (before date filtering) for first KPI card
            if report_type == "PEMEDES Complaints Only":
                all_time_total = len(df_base)  # All PEMEDES complaints ever
            elif report_type == "NTC Complaints Only":
                all_time_total = len(df_base)  # All NTC complaints ever
            else:  # Total (All Complaints)
                all_time_total = len(df_base)  # All complaints ever
            
            # Calculate period-specific metrics on filtered data for second KPI card
            period_total = len(df_filtered)
            
            if report_type == "PEMEDES Complaints Only":
                ntc_custom_count = 0  # Not applicable for PEMEDES report
                pemedes_custom_count = period_total  # Period complaints are PEMEDES
            elif report_type == "NTC Complaints Only":
                ntc_custom_count = period_total  # Period complaints are NTC
                pemedes_custom_count = 0  # Not applicable for NTC report
            else:  # Total (All Complaints)
                # NTC Calculation - Use ONLY Telco Internet Issues for exact match
                ntc_custom_count = 0
                if 'Complaint Category' in df_filtered.columns:
                    ntc_mask_custom = (df_filtered['Complaint Category'].astype(str).str.strip().str.upper() == "TELCO INTERNET ISSUES")
                    ntc_custom_count = len(df_filtered[ntc_mask_custom])
                
                # PEMEDES Calculation
                pemedes_custom_count = 0
                if 'Complaint Category' in df_filtered.columns:
                    pemedes_mask_custom = (df_filtered['Complaint Category'].astype(str).str.strip().str.upper() == "DELIVERY CONCERNS (SP)")
                    pemedes_custom_count = len(df_filtered[pemedes_mask_custom])

            # Store metrics for export
            metrics = {
                'total': period_total,  # Use period total for consistency with report generation
                'ntc': ntc_custom_count,
                'pemedes': pemedes_custom_count,
                'start_date': start_date,
                'end_date': end_date
            }

            # Display Metrics - 2 KPI Cards Only
            
            if report_type == "PEMEDES Complaints Only":
                # Show PEMEDES all-time vs period metrics
                c_col1, c_col2 = st.columns(2)
                c_col1.metric("Total Complaints", f"{all_time_total:,}", help=f"All-time delivery complaints in the database")
                period_label = coverage_period.lower() if coverage_period else "selected"
                c_col2.metric("PEMEDES Period", f"{period_total:,}", help=f"Delivery concerns during {period_label} period ({start_date.strftime('%b %d, %Y')} to {end_date.strftime('%b %d, %Y')})")
            elif report_type == "NTC Complaints Only":
                # Show NTC all-time vs period metrics
                c_col1, c_col2 = st.columns(2)
                c_col1.metric("Total Complaints", f"{all_time_total:,}", help=f"All-time telecommunications complaints in the database")
                period_label = coverage_period.lower() if coverage_period else "selected"
                c_col2.metric("NTC Period", f"{period_total:,}", help=f"Telecom issues during {period_label} period ({start_date.strftime('%b %d, %Y')} to {end_date.strftime('%b %d, %Y')})")
            else:  # Total (All Complaints)
                # Show all-time vs period comprehensive metrics
                c_col1, c_col2 = st.columns(2)
                c_col1.metric("Total Complaints", f"{all_time_total:,}", help=f"All-time complaints in the database")
                period_label = coverage_period.lower() if coverage_period else "selected"
                c_col2.metric("Period Total", f"{period_total:,}", help=f"All complaints during {period_label} period ({start_date.strftime('%b %d, %Y')} to {end_date.strftime('%b %d, %Y')})")
            
            # Use filtered data for the report
            df = df_filtered
        else:
            st.warning("No valid dates found in data.")
    else:
        st.warning("Date Received column missing. Cannot filter by date.")

    # Data alignment confirmation
    total_records = len(df)
    date_range_info = ""
    if 'Date Received' in df.columns:
        valid_dates = df['Date Received'].dropna()
        if len(valid_dates) > 0:
            min_date = valid_dates.min()
            max_date = valid_dates.max()
            date_range_info = f" | Data Range: {min_date.strftime('%b %Y')} - {max_date.strftime('%b %Y')}"

    # Initialize Vertex AI
    is_init, error_msg = init_vertex_ai()

    if not is_init:
        st.warning("Vertex AI not initialized. AI features are currently disabled.")
        with st.expander("Error Details"):
            if error_msg:
                st.caption(f"Error: {error_msg}")
                st.caption("Please ensure you have 'Vertex AI User' role and the API is enabled in Google Cloud.")

    # Get Top Issues (always compute to show preview)
    top_issues = get_top_issues(df)

    if not top_issues:
        st.warning("Insufficient data to identify top issues. Please ensure your data has 'Complaint Category' or 'Complaint Nature' columns.")
        return

    # Generation Button - Centered and prominent
    st.markdown("---")
    col_spacer1, col_btn, col_spacer2 = st.columns([1, 2, 1])

    # Use unique session state keys for each report type and coverage period
    report_key = f"{report_type.replace(' ', '_').replace('(', '').replace(')', '')}_{coverage_period}"
    
    with col_btn:
        if st.session_state.get(f'report_generated_{report_key}', False):
            generate_button = st.button(
                f"🔄 Regenerate {report_type} Action Plan",
                type="secondary",
                use_container_width=True,
                help=f"Re-analyze top complaints and regenerate strategic action plans for {report_type.lower()}"
            )
        else:
            generate_button = st.button(
                f"Generate {report_type} Action Plan",
                type="primary",
                use_container_width=True,
                help=f"Analyze top complaints and generate strategic action plans for {report_type.lower()}"
            )
    
    # Add padding below the Generate button
    st.markdown("<br>", unsafe_allow_html=True)

    # Handle generation
    if generate_button or (f'weekly_action_plan_{report_key}' in st.session_state and st.session_state.get(f'report_generated_{report_key}', False)):
        if generate_button:
            # Center the spinner message for better UX
            spinner_col1, spinner_col2, spinner_col3 = st.columns([1, 2, 1])
            with spinner_col2:
                with st.spinner(f"Analyzing {report_type.lower()} data and generating strategic action plans..."):
                    if is_init:
                        action_plan_data = generate_ai_action_plan(top_issues, df)  # Pass df for SP analysis
                        # Generate Executive Summary
                        summary_data = generate_executive_summary(action_plan_data)
                    else:
                        # Fallback if no AI - generate concise action plans with minimal data-driven remarks
                        action_plan_data = []
                        for i in top_issues:
                            unit_code, unit_name, org_type = categorize_issue_to_unit(i['name'], i['type'])

                            # Generate minimal data-driven remarks (non-AI scenario)
                            count = i.get('count', 0)
                            remarks = f"Affects {count:,} complaints. Assigned to {unit_code} for resolution."

                            action_plan_data.append({
                                "issue": i['name'],
                                "action_plan": f"Coordinate with {unit_code} to address complaints",
                                "unit": unit_code,
                                "remarks": remarks
                            })
                        summary_data = {"main_summary": "AI features unavailable.", "org_summaries": {}}

            st.session_state[f'weekly_action_plan_{report_key}'] = action_plan_data
            st.session_state[f'executive_summary_{report_key}'] = summary_data
            st.session_state[f'report_generated_{report_key}'] = True
            st.session_state[f'report_timestamp_{report_key}'] = datetime.now()
            st.success(f"{report_type} action plan generated successfully.")
            
            # Auto-scroll to results for better UX
            st.markdown('<script>window.scrollTo(0, document.body.scrollHeight);</script>', unsafe_allow_html=True)

    # Display Report Content if Generated
    if st.session_state.get(f'report_generated_{report_key}', False):
        
        # Executive Summary Section
        st.markdown(f"### I. Report Overview - {report_type}")
        if f'executive_summary_{report_key}' in st.session_state:
            st.markdown(f"""
            <div style="background-color: #f8fafc; padding: 1.5rem; border-radius: 8px; border-left: 4px solid #3b82f6; margin-bottom: 1rem;">
                <p style="font-size: 1.05rem; line-height: 1.6; color: #1f2937; margin: 0;">
                    {st.session_state[f'executive_summary_{report_key}'].get('main_summary', 'Summary unavailable.')}
                </p>
            </div>
            """, unsafe_allow_html=True)
        
        # Top Issues Section
        st.markdown("### II. Top 5 Priority Issues")
        
        # Calculate total complaints in top 5
        total_top5 = sum([issue['count'] for issue in top_issues])
        coverage_pct = (total_top5 / total_records * 100) if total_records > 0 else 0

        st.markdown(f"""
        <div class="info-card" style="background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%); border-left: 4px solid #f59e0b;">
            <strong>Impact Analysis:</strong> The top 5 priority issues represent <strong>{total_top5:,} complaints ({coverage_pct:.1f}%)</strong> of the total dataset.
            Addressing these priorities will have the most significant impact on reducing overall complaint volume.
        </div>
        """, unsafe_allow_html=True)

        # Enrich issues with unit recommendations
        preview_data = []
        for idx, issue in enumerate(top_issues, 1):
            unit_code, unit_name, org_type = categorize_issue_to_unit(issue['name'], issue['type'])
            preview_data.append({
                "#": idx,
                "Issue": issue['name'],
                "Source": issue['type'],
                "Count": issue['count'],
                "Recommended Unit": unit_code,
                "Organization": org_type
            })

        preview_df = pd.DataFrame(preview_data)
        st.dataframe(
            preview_df,
            column_config={
                "#": st.column_config.NumberColumn("Rank", width="small"),
                "Issue": st.column_config.TextColumn("Issue Description", width="large"),
                "Source": st.column_config.TextColumn("Source", width="small"),
                "Count": st.column_config.NumberColumn("Complaints", format="%d", width="small"),
                "Recommended Unit": st.column_config.TextColumn("Assigned Unit", width="medium"),
                "Organization": st.column_config.TextColumn("Organization Type", width="medium")
            },
            hide_index=True,
            use_container_width=True
        )
        st.caption("**Note:** Source indicates whether the issue is from Category or Nature field. Assigned Unit is auto-categorized based on DICT organizational structure.")

        st.markdown("---")

        # DICT Unit Analysis - Calculation for Export Only
        dict_unit_counts = None
        if 'DICT UNIT' in df.columns:
            valid_units = df['DICT UNIT'].dropna()
            valid_units = valid_units[valid_units != '']
            
            # Exclude NTC from DICT Unit analysis
            valid_units = valid_units[~valid_units.astype(str).str.upper().isin(['NTC', 'NATIONAL TELECOMMUNICATIONS COMMISSION'])]
            
            if len(valid_units) > 0:
                dict_unit_counts = valid_units.value_counts().head(10)

                st.success("Action plan generated successfully.")

        # Display the action plan
        if f'weekly_action_plan_{report_key}' in st.session_state:
            plans = st.session_state[f'weekly_action_plan_{report_key}']

            # Validate plans data
            if not plans or len(plans) == 0:
                st.warning("No action plans were generated. Please try again.")
                return

            # Convert to DataFrame and add editable columns
            report_df = pd.DataFrame(plans)

            # Validate dataframe is not empty
            if report_df.empty:
                st.warning("Action plan data is empty. Please regenerate the report.")
                return

            # Add resolution column if it doesn't exist (internal name stays 'resolution')
            if 'resolution' not in report_df.columns:
                report_df['resolution'] = ''

            # Rename columns for better display
            display_df = report_df.rename(columns={
                'issue': 'Top Issue',
                'action_plan': 'Action Plan',
                'unit': 'Assigned Unit',
                'remarks': 'Remarks',
                'resolution': 'Action Taken by the Unit'
            })

            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("---")
            st.markdown("### II. Strategic Action Plan Details")
            st.caption(f"Generated {len(report_df)} strategic recommendations based on analysis of top complaint patterns")

            # Info box about editing
            st.info("💡 **Fully Editable Table!** All fields can be edited. Make your changes, then scroll to the bottom and click 'Save All Changes' to apply.")

            # Initialize edited_action_plan on first load only (per report type)
            if f'edited_action_plan_{report_key}' not in st.session_state:
                st.session_state[f'edited_action_plan_{report_key}'] = report_df.copy()
                st.session_state[f'data_changed_{report_key}'] = True

            # Use saved data for display (or original if not yet saved)
            display_source_df = st.session_state[f'edited_action_plan_{report_key}'].copy()

            # Ensure resolution column exists
            if 'resolution' not in display_source_df.columns:
                display_source_df['resolution'] = ''

            # Rename columns for display
            display_df_for_editor = display_source_df.rename(columns={
                'issue': 'Top Issue',
                'action_plan': 'Action Plan',
                'unit': 'Assigned Unit',
                'remarks': 'Remarks',
                'resolution': 'Action Taken by the Unit'
            })

            # Editable data editor - ALL fields are editable with proper wrapping
            edited_df = st.data_editor(
                display_df_for_editor,
                column_config={
                    "Top Issue": st.column_config.TextColumn(
                        "Top Issue",
                        width=180,
                        disabled=False,
                        help="Click to edit issue name"
                    ),
                    "Action Plan": st.column_config.TextColumn(
                        "Action Plan",
                        width=400,
                        disabled=False,
                        help="Click to edit the action plan"
                    ),
                    "Assigned Unit": st.column_config.TextColumn(
                        "Assigned Unit",
                        width=100,
                        disabled=False,
                        help="Click to change assigned unit"
                    ),
                    "Remarks": st.column_config.TextColumn(
                        "Remarks",
                        width=250,
                        disabled=False,
                        help="Click to add or edit remarks"
                    ),
                    "Action Taken by the Unit": st.column_config.TextColumn(
                        "Action Taken by the Unit",
                        width=200,
                        disabled=False,
                        help="Click to update actions taken"
                    )
                },
                hide_index=True,
                use_container_width=True,
                num_rows="fixed",
                key="action_plan_editor"
            )

            # Store edited main table temporarily (local variable only - no state change)
            temp_edited_main_df = edited_df

            # Service Provider Breakdown for PRD and NTC issues
            st.markdown("---")
            st.markdown("### III. Service Provider Analysis")
            st.caption("Detailed breakdown for Delivery Concerns and Telecommunications Issues")

            # Use edited dataframe for further processing
            current_report_df = st.session_state[f'edited_action_plan_{report_key}']

            # Initialize service provider breakdowns in session state if not exists (per report type)
            if f'sp_breakdowns_{report_key}' not in st.session_state:
                st.session_state[f'sp_breakdowns_{report_key}'] = {}

            # Check which issues need SP breakdown
            issues_with_breakdown = []
            temp_sp_edits_local = {}  # Local dictionary to collect SP edits (no state changes)

            for idx, row in current_report_df.iterrows():
                unit = row['unit']
                issue_name = row['issue']

                # Find matching issue from top_issues
                matching_issue = next((i for i in top_issues if i['name'] == issue_name), None)

                if matching_issue and unit in UNITS_REQUIRING_SP_BREAKDOWN:
                    # Use cached breakdown if exists, otherwise fetch new
                    sp_key = f"{issue_name}_{unit}"
                    if sp_key not in st.session_state[f'sp_breakdowns_{report_key}']:
                        sp_breakdown = get_service_provider_breakdown(df, issue_name, matching_issue['type'])
                        if sp_breakdown:
                            st.session_state[f'sp_breakdowns_{report_key}'][sp_key] = sp_breakdown
                    else:
                        sp_breakdown = st.session_state[f'sp_breakdowns_{report_key}'][sp_key]

                    if sp_breakdown:
                        issues_with_breakdown.append({
                            "issue": issue_name,
                            "unit": unit,
                            "unit_label": UNITS_REQUIRING_SP_BREAKDOWN[unit],
                            "total_count": matching_issue['count'],
                            "breakdown": sp_breakdown,
                            "sp_key": sp_key
                        })

            if issues_with_breakdown:
                st.info("💡 Service Provider tables are also editable. Edit as needed, then use 'Save All Changes' button at the bottom.")

                for idx, item in enumerate(issues_with_breakdown):
                    with st.expander(f"{item['issue']} ({item['unit']}) - {item['total_count']} complaints", expanded=True):
                        # Create editable breakdown table from session state
                        sp_df = pd.DataFrame(item['breakdown'])

                        edited_sp_df = st.data_editor(
                            sp_df,
                            column_config={
                                "provider": st.column_config.TextColumn(
                                    "Service Provider",
                                    width=300,
                                    disabled=False,
                                    help="Click to edit provider name"
                                ),
                                "count": st.column_config.NumberColumn(
                                    "Complaints",
                                    format="%d",
                                    width=120,
                                    disabled=False,
                                    help="Click to edit count"
                                ),
                                "percentage": st.column_config.NumberColumn(
                                    "Percentage",
                                    format="%.1f%%",
                                    width=120,
                                    disabled=False,
                                    help="Click to edit percentage"
                                )
                            },
                            hide_index=True,
                            use_container_width=True,
                            num_rows="fixed",
                            key=f"sp_breakdown_{idx}",
                            disabled=False
                        )

                        # Store edits in local dictionary (NO session state update - prevents rerun!)
                        new_sp_breakdown = edited_sp_df.to_dict('records')
                        item['breakdown'] = new_sp_breakdown
                        temp_sp_edits_local[item['sp_key']] = new_sp_breakdown

                        # Summary stats
                        st.caption(f"Top Provider: **{item['breakdown'][0]['provider']}** with {item['breakdown'][0]['count']} complaints ({item['breakdown'][0]['percentage']}%)")
                        st.caption(f"Total Providers Identified: {len(item['breakdown'])}")

                        # Add concise explanation
                        st.markdown("---")
                        st.markdown("**Analysis and Recommendations:**")

                        # Calculate insights
                        top_provider = item['breakdown'][0]
                        top_provider_pct = top_provider['percentage']
                        num_providers = len(item['breakdown'])

                        # Generate contextual explanation based on the data
                        if top_provider_pct > 50:
                            concentration = "highly concentrated"
                            recommendation = f"Focus immediate attention on **{top_provider['provider']}** as they account for the majority of issues. Consider escalating to their management team."
                        elif top_provider_pct > 30:
                            concentration = "moderately concentrated"
                            recommendation = f"Prioritize **{top_provider['provider']}** while monitoring other providers. A targeted intervention could significantly reduce complaints."
                        else:
                            concentration = "distributed across multiple providers"
                            recommendation = f"Issues are spread across {num_providers} providers. Consider a broader systemic approach rather than targeting individual providers."

                        st.write(f"Complaint distribution for this issue is **{concentration}**, with the leading provider accounting for **{top_provider_pct:.1f}%** of all complaints in this category. {recommendation}")
            else:
                st.info("No Delivery Concerns or Telecommunications Issues identified in the top 5 priority complaints requiring detailed service provider analysis.")

            st.markdown("---")

            # Save processing will be handled at the end after all sections

            # Export logic moved to end for better user experience

            # Summary Metrics
            st.markdown("---")
            st.markdown(f"### IV. Executive Summary - {report_type}")
            
            if f'executive_summary_{report_key}' in st.session_state:
                st.markdown(f"""
                <div style="background-color: #f8fafc; padding: 1.5rem; border-radius: 8px; border-left: 4px solid #3b82f6; margin-bottom: 1rem;">
                    <p style="font-size: 1.05rem; line-height: 1.6; color: #1f2937; margin: 0;">
                        {st.session_state[f'executive_summary_{report_key}'].get('main_summary', 'Summary unavailable.')}
                    </p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.info(f"Executive summary not available for {report_type.lower()}.")

            # Unit Distribution Analysis with Specific Details
            st.markdown("---")
            st.markdown("### V. Unit Assignment and Service Provider Details")
            st.caption("Agencies and service providers requiring intervention")

            # Use edited dataframe for display (get from session state)
            export_df = st.session_state.get(f'edited_action_plan_{report_key}', st.session_state.get(f'weekly_action_plan_{report_key}', pd.DataFrame()))
            
            # Build detailed breakdown with service providers
            unit_details = []
            for idx, row in export_df.iterrows():
                unit = row['unit']
                issue_name = row['issue']

                # Skip N/A units since they are not valid DICT units
                if unit.upper() in ['N/A', 'N.A', 'NA', 'NOT APPLICABLE', 'NONE', 'N./A']:
                    continue

                # Get matching issue from top_issues
                matching_issue = next((i for i in top_issues if i['name'] == issue_name), None)

                # Get top service provider if applicable
                top_provider = None
                provider_count = 0
                if matching_issue and unit in UNITS_REQUIRING_SP_BREAKDOWN:
                    sp_breakdown = get_service_provider_breakdown(df, issue_name, matching_issue['type'])
                    if sp_breakdown and len(sp_breakdown) > 0:
                        top_provider = sp_breakdown[0]['provider']
                        provider_count = sp_breakdown[0]['count']

                # Get full unit name
                unit_full_name = DICT_UNIT_MAPPING.get(unit, {}).get("name", unit)

                # Categorize
                if unit in DELIVERY_UNITS:
                    category = "Delivery Unit (DICT)"
                elif unit in ATTACHED_AGENCIES:
                    category = "Attached Agency"
                elif unit in OTHER_AGENCIES:
                    category = "External Agency"
                else:
                    category = "Unclassified"

                unit_details.append({
                    "Unit Code": unit,
                    "Unit Name": unit_full_name,
                    "Category": category,
                    "Issue": issue_name,
                    "Top Service Provider": top_provider if top_provider else "N/A",
                    "SP Complaints": provider_count if top_provider else 0,
                    "Total Complaints": matching_issue['count'] if matching_issue else 0
                })

            # Display as comprehensive table
            details_df = pd.DataFrame(unit_details)
            st.dataframe(
                details_df,
                column_config={
                    "Unit Code": st.column_config.TextColumn("Unit", width="small"),
                    "Unit Name": st.column_config.TextColumn("Agency/Unit Name", width="large"),
                    "Category": st.column_config.TextColumn("Type", width="medium"),
                    "Issue": st.column_config.TextColumn("Issue Assigned", width="large"),
                    "Top Service Provider": st.column_config.TextColumn("Top Service Provider", width="medium"),
                    "SP Complaints": st.column_config.NumberColumn("SP Count", format="%d", width="small"),
                    "Total Complaints": st.column_config.NumberColumn("Total", format="%d", width="small")
                },
                hide_index=True,
                use_container_width=True
            )

            # Summary by category
            st.markdown(f"#### A. Summary by Organization Type - {report_type}")
            
            if f'executive_summary_{report_key}' in st.session_state:
                org_summaries = st.session_state[f'executive_summary_{report_key}'].get('org_summaries', {})
                
                # Display summaries for each category if available
                # Map the keys returned by AI to our display titles
                # AI Prompt asked for: "Delivery Unit (DICT Internal)", "Attached Agency", "External Agency"
                
                categories = [
                    ("DICT Delivery Units", ["Delivery Unit", "DICT Internal", "Delivery Unit (DICT Internal)"]),
                    ("Attached Agencies", ["Attached Agency", "Attached Agencies"]),
                    ("External Agencies", ["External Agency", "External Agencies"])
                ]
                
                has_summary = False
                for title, keys in categories:
                    summary_text = None
                    # Find matching summary
                    for k, v in org_summaries.items():
                        if any(key_part.lower() in k.lower() for key_part in keys):
                            summary_text = v
                            break
                    
                    if summary_text:
                        has_summary = True
                        st.markdown(f"**{title}**")
                        st.markdown(f"""
                        <div style="background-color: #f9fafb; padding: 1rem; border-radius: 6px; border: 1px solid #e5e7eb; margin-bottom: 0.75rem;">
                            <p style="font-size: 0.95rem; color: #374151; margin: 0;">{summary_text}</p>
                        </div>
                        """, unsafe_allow_html=True)
                
                if not has_summary:
                    st.info("No specific organization summaries generated.")
            else:
                st.info("Organization summaries not available.")

            # Validation status
            unclassified = [d for d in unit_details if d['Category'] == "Unclassified"]
            
            # Count N/A units that were filtered out
            na_count = 0
            for idx, row in export_df.iterrows():
                unit = row['unit']
                if unit.upper() in ['N/A', 'N.A', 'NA', 'NOT APPLICABLE', 'NONE', 'N./A']:
                    na_count += 1
            
            if unclassified:
                st.warning(f"{len(unclassified)} issue(s) could not be categorized properly")
                for item in unclassified:
                    st.caption(f"• {item['Unit Code']} - {item['Issue']}")
            else:
                st.success("All valid issues successfully categorized to appropriate units and agencies")
            
            if na_count > 0:
                st.info(f"ℹ️ {na_count} issue(s) with N/A units were excluded from analysis (not applicable for unit assignment)")

            # Save All Changes Section (moved to end for better UX)
            st.markdown("---")
            st.markdown("### VI. Save All Changes")
            st.caption("Apply your edits and prepare export files")
            
            col_save1, col_save2, col_save3 = st.columns([1, 1, 1])
            with col_save2:
                save_button = st.button(
                    "💾 Save All Changes",
                    type="primary",
                    use_container_width=True,
                    help="Save edits from main table and service provider tables, then regenerate export files"
                )

        # Process save button (outside the form but still in the function)
        if save_button:
            # Convert edited main table data back to internal column names
            temp_edited_main = temp_edited_main_df.rename(columns={
                'Top Issue': 'issue',
                'Action Plan': 'action_plan',
                'Assigned Unit': 'unit',
                'Remarks': 'remarks',
                'Action Taken by the Unit': 'resolution'
            })

            # Update session state with main table edits (per report type)
            st.session_state[f'edited_action_plan_{report_key}'] = temp_edited_main

            # Apply all SP edits to permanent storage (per report type)
            for sp_key, sp_data in temp_sp_edits_local.items():
                st.session_state[f'sp_breakdowns_{report_key}'][sp_key] = sp_data

            # Mark that export cache needs regeneration (per report type)
            st.session_state[f'data_changed_{report_key}'] = True

            st.success(f"✅ All changes for {report_type} saved successfully! Export files will be updated.")
            st.rerun()

        # Export Options Section (moved to end for better UX)
        st.markdown("---")
        st.markdown("### VII. Export Options")
        st.caption("Download the strategic action plan in your preferred format (includes your edits)")

        # Use edited dataframe for exports (with safety check)
        if f'edited_action_plan_{report_key}' not in st.session_state:
            st.info(f"Click 'Save All Changes' above to prepare {report_type.lower()} exports.")
        else:
            export_df = st.session_state[f'edited_action_plan_{report_key}']

            # Cache the export data as bytes to prevent regeneration on download (per report type)
            if f'cached_pdf_bytes_{report_key}' not in st.session_state or st.session_state.get(f'data_changed_{report_key}', True):
                try:
                    pdf_buffer = export_to_pdf(export_df, top_issues, issues_with_breakdown, dict_unit_counts, st.session_state.get(f'executive_summary_{report_key}'), metrics, report_type)
                    st.session_state[f'cached_pdf_bytes_{report_key}'] = pdf_buffer.getvalue()  # Store as bytes
                    st.session_state[f'pdf_error_{report_key}'] = None
                except Exception as e:
                    st.session_state[f'cached_pdf_bytes_{report_key}'] = None
                    st.session_state[f'pdf_error_{report_key}'] = str(e)

            if f'cached_word_bytes_{report_key}' not in st.session_state or st.session_state.get(f'data_changed_{report_key}', True):
                try:
                    word_buffer = export_to_word(export_df, top_issues, issues_with_breakdown, dict_unit_counts, st.session_state.get(f'executive_summary_{report_key}'), metrics, report_type)
                    st.session_state[f'cached_word_bytes_{report_key}'] = word_buffer.getvalue()  # Store as bytes
                    st.session_state[f'word_error_{report_key}'] = None
                except Exception as e:
                    st.session_state[f'cached_word_bytes_{report_key}'] = None
                    st.session_state[f'word_error_{report_key}'] = str(e)

            if f'cached_csv_string_{report_key}' not in st.session_state or st.session_state.get(f'data_changed_{report_key}', True):
                st.session_state[f'cached_csv_string_{report_key}'] = export_df.to_csv(index=False)

            # Mark data as cached (per report type)
            st.session_state[f'data_changed_{report_key}'] = False

            col_dl1, col_dl2, col_dl3 = st.columns(3)

            # Generate filename suffix based on report type
            file_suffix = ""
            if report_type == "PEMEDES Complaints Only":
                file_suffix = "_PEMEDES"
            elif report_type == "NTC Complaints Only":
                file_suffix = "_NTC"
            else:
                file_suffix = "_Total"
                
            with col_dl1:
                # PDF Download
                if st.session_state.get(f'cached_pdf_bytes_{report_key}'):
                    st.download_button(
                        label="📄 PDF Document",
                        data=st.session_state[f'cached_pdf_bytes_{report_key}'],
                        file_name=f"DICT_AI_Action_Plan{file_suffix}_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        help=f"Download formatted PDF report for {report_type.lower()}",
                        key=f"download_pdf_btn_{report_key}"
                    )
                else:
                    st.error(f"PDF Export Error: {st.session_state.get(f'pdf_error_{report_key}', 'Unknown error')}")

            with col_dl2:
                # Word Download
                if st.session_state.get(f'cached_word_bytes_{report_key}'):
                    st.download_button(
                        label="📝 Word Document",
                        data=st.session_state[f'cached_word_bytes_{report_key}'],
                        file_name=f"DICT_AI_Action_Plan{file_suffix}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        help=f"Download editable Word document for {report_type.lower()}",
                        key=f"download_word_btn_{report_key}"
                    )
                else:
                    st.error(f"Word Export Error: {st.session_state.get(f'word_error_{report_key}', 'Unknown error')}")

            with col_dl3:
                # CSV Download
                st.download_button(
                    label="📊 CSV Spreadsheet",
                    data=st.session_state[f'cached_csv_string_{report_key}'],
                    file_name=f"DICT_AI_Action_Plan{file_suffix}_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    use_container_width=True,
                    help=f"Download CSV data for {report_type.lower()}",
                    key=f"download_csv_btn_{report_key}"
                )
            
            st.caption("**Note:** Export files include all saved changes from the editable tables above.")
